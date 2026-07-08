#!/usr/bin/env python3
"""
生成「移动应用开发实践」课程设计报告 DOCX — 严格按模板格式
项目：朝夕 — 生活秩序管理助手
个人负责：鸿蒙前端UI设计与框架搭建 + Python Flask中转服务器

格式规范（来自课程设计报告模板）：
- 页边距：左/上 2.5cm，右/下 2.0cm
- 正文：小四(12pt)，宋体+Times New Roman，固定值22磅行距，首行缩进2字符，两端对齐
- 一级标题(Heading 1)：小三(15pt)黑体，居中，段前段后1行(22pt)
- 二级标题(Heading 2)：四号(14pt)黑体，左顶格
- 三级标题(Heading 3)：小四(12pt)黑体，左顶格
- 图题/表题：五号(10.5pt)，居中，无缩进
- 封面无页码，目录页罗马数字，正文阿拉伯数字从1开始
- 每章另起一页（通过 Heading 1 样式设置 pageBreakBefore）
- 目录用 TOC 域自动生成
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ============================================================
# 配置
# ============================================================
STUDENT_NAME = "温志远"
STUDENT_ID   = "xxxxxxxxxx"
COLLEGE      = "计算机科学与通信工程学院"
CLASS_NAME   = "通信工程 2022级01班"
ADVISORS     = "单田华 朱轶"
COURSE_START = "2026.06.22"
COURSE_END   = "2026.07.05"
OUTPUT_FILE  = f"/Users/wenzhiyuan/Desktop/移动应用开发实践课程设计报告-{STUDENT_ID}{STUDENT_NAME}.docx"

# ============================================================
# XML 工具
# ============================================================
def _set_exact_line_spacing(paragraph, pt_val):
    """固定值行距 - 用于段落"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.insert(0, spacing)
    spacing.set(qn('w:line'), str(int(pt_val * 20)))
    spacing.set(qn('w:lineRule'), 'exact')

def _set_style_line_spacing(style, pt_val):
    """固定值行距 - 用于样式（修改样式内的 w:pPr）"""
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style.element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(pt_val * 20)))
    spacing.set(qn('w:lineRule'), 'exact')

def _run_font(run, western='Times New Roman', east_asian='宋体', size=Pt(12), bold=False):
    """同时设置西文和中文字体"""
    run.font.size = size
    run.bold = bold
    run.font.name = western
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), western)
    rFonts.set(qn('w:hAnsi'), western)
    rFonts.set(qn('w:eastAsia'), east_asian)
    rFonts.set(qn('w:cs'), western)

def _add_run(paragraph, text, western='Times New Roman', east_asian='宋体',
             size=Pt(12), bold=False):
    run = paragraph.add_run(text)
    _run_font(run, western, east_asian, size, bold)
    return run

def _configure_style_font(style, western='Times New Roman', east_asian='宋体',
                          size=Pt(12), bold=False):
    """配置样式级别的字体"""
    style.font.size = size
    style.font.bold = bold
    style.font.name = western
    style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，覆盖 Word 默认蓝色标题
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), western)
    rFonts.set(qn('w:hAnsi'), western)
    rFonts.set(qn('w:eastAsia'), east_asian)
    rFonts.set(qn('w:cs'), western)
    # 确保 w:color 写入 XML
    color_el = rPr.find(qn('w:color'))
    if color_el is None:
        color_el = OxmlElement('w:color')
        rPr.append(color_el)
    color_el.set(qn('w:val'), '000000')

def _apply_body_defaults(paragraph):
    """正文默认格式"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(24)
    _set_exact_line_spacing(paragraph, 22)

# ============================================================
# 内容段落函数（不设样式，仅排版）
# ============================================================
def body(doc, text):
    p = doc.add_paragraph()
    _apply_body_defaults(p)
    _add_run(p, text)
    return p

def bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Pt(0)
    _set_exact_line_spacing(p, 22)
    _add_run(p, '● ' + text)
    return p

def table_caption(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    _set_exact_line_spacing(p, 18)
    _add_run(p, text, size=Pt(10.5))
    return p

def create_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_exact_line_spacing(p, 18)
        run = p.add_run(h)
        _run_font(run, east_asian='宋体', size=Pt(10.5), bold=True)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9" w:val="clear"/>')
        c._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            _set_exact_line_spacing(p, 18)
            run = p.add_run(str(val))
            _run_font(run, east_asian='宋体', size=Pt(10.5))
    return table

# ============================================================
# 标题函数 — 使用内置样式 Heading 1/2/3
# ============================================================
def heading1(doc, text):
    """一级标题：Heading 1 样式"""
    return doc.add_paragraph(text, style='Heading 1')

def heading2(doc, text):
    """二级标题：Heading 2 样式"""
    return doc.add_paragraph(text, style='Heading 2')

def heading3(doc, text):
    """三级标题：Heading 3 样式"""
    return doc.add_paragraph(text, style='Heading 3')

# ============================================================
# TOC 域
# ============================================================
def insert_toc_field(doc):
    """插入 Word TOC 域，打开后右键更新域即可自动生成目录"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)

    # w:fldChar begin
    run_begin = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_begin._element.append(fldChar_begin)

    # w:instrText
    run_instr = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z '
    run_instr._element.append(instrText)

    # w:fldChar separate
    run_sep = p.add_run()
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run_sep._element.append(fldChar_sep)

    # 占位文字
    run_text = p.add_run('（请在 Word 中右键此处 → 更新域，自动生成目录）')
    _run_font(run_text, size=Pt(12))

    # w:fldChar end
    run_end = p.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_end._element.append(fldChar_end)

    return p

# ============================================================
# 节属性配置
# ============================================================
def _set_sect_pr(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

def _configure_page_numbers(section, start_num, num_fmt='decimal'):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.insert(1, pgNumType)
    pgNumType.set(qn('w:fmt'), num_fmt)
    pgNumType.set(qn('w:start'), str(start_num))

def _add_title_pg(section):
    """首页不同（封面不显示页码）"""
    sectPr = section._sectPr
    if sectPr.find(qn('w:titlePg')) is None:
        sectPr.insert(0, OxmlElement('w:titlePg'))

def _remove_title_pg(section):
    tp = section._sectPr.find(qn('w:titlePg'))
    if tp is not None:
        section._sectPr.remove(tp)

# ============================================================
# 主生成
# ============================================================
def generate():
    doc = Document()

    # ============ 配置样式 ============
    # Normal 样式基准
    style_normal = doc.styles['Normal']
    _configure_style_font(style_normal, size=Pt(12))
    pf_n = style_normal.paragraph_format
    pf_n.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading 1: 小三(15pt)黑体，居中，段前段后22pt，每章另起一页
    h1 = doc.styles['Heading 1']
    _configure_style_font(h1, east_asian='黑体', size=Pt(15), bold=True)
    h1pf = h1.paragraph_format
    h1pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1pf.space_before = Pt(22)
    h1pf.space_after = Pt(22)
    h1pf.first_line_indent = Pt(0)
    h1pf.page_break_before = True   # 每章另起一页
    _set_style_line_spacing(h1, 22)

    # Heading 2: 四号(14pt)黑体，左顶格
    h2 = doc.styles['Heading 2']
    _configure_style_font(h2, east_asian='黑体', size=Pt(14), bold=True)
    h2pf = h2.paragraph_format
    h2pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2pf.space_before = Pt(10)
    h2pf.space_after = Pt(6)
    h2pf.first_line_indent = Pt(0)
    _set_style_line_spacing(h2, 22)

    # Heading 3: 小四(12pt)黑体，左顶格
    h3 = doc.styles['Heading 3']
    _configure_style_font(h3, east_asian='黑体', size=Pt(12), bold=True)
    h3pf = h3.paragraph_format
    h3pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3pf.space_before = Pt(6)
    h3pf.space_after = Pt(4)
    h3pf.first_line_indent = Pt(0)
    _set_style_line_spacing(h3, 22)

    # ============ Section 1: 封面 ============
    sec1 = doc.sections[0]
    _set_sect_pr(sec1)
    _add_title_pg(sec1)

    for _ in range(4):
        p = doc.add_paragraph(); _apply_body_defaults(p); _add_run(p, '', size=Pt(12))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, '江  苏  大  学', east_asian='黑体', size=Pt(26), bold=True)

    for _ in range(3):
        p = doc.add_paragraph(); _apply_body_defaults(p); _add_run(p, '', size=Pt(12))

    for t in ['移动应用开发实践', '课程设计报告']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, t, east_asian='黑体', size=Pt(22), bold=True)

    for _ in range(4):
        p = doc.add_paragraph(); _apply_body_defaults(p); _add_run(p, '', size=Pt(12))

    for label in [
        '学院名称：' + COLLEGE, '专业班级：' + CLASS_NAME,
        '学生姓名：' + STUDENT_NAME, '学生学号：' + STUDENT_ID,
        '指导教师姓名：' + ADVISORS, '课程设计时间：' + f'{COURSE_START}-{COURSE_END}',
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, label, size=Pt(14))

    for _ in range(3):
        p = doc.add_paragraph(); _apply_body_defaults(p); _add_run(p, '', size=Pt(12))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, '2026年07月', size=Pt(14))

    # ============ Section 2: 目录页（罗马页码） ============
    doc.add_section()
    sec2 = doc.sections[-1]
    _set_sect_pr(sec2)
    _configure_page_numbers(sec2, 1, num_fmt='upperRoman')

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(22)
    p.paragraph_format.first_line_indent = Pt(0)
    _add_run(p, '目  录', east_asian='黑体', size=Pt(18), bold=True)

    # 插入 TOC 域（打开 Word 后右键更新域）
    insert_toc_field(doc)

    # ============ Section 3: 正文 ============
    doc.add_section()
    sec3 = doc.sections[-1]
    _set_sect_pr(sec3)
    _configure_page_numbers(sec3, 1, num_fmt='decimal')
    _remove_title_pg(sec3)

    header = sec3.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(hp, '移动应用开发实践课程设计报告', size=Pt(10.5))

    # ============================================================
    # 第一章
    # ============================================================
    heading1(doc, '第一章  引言')

    heading2(doc, '1.1  课程设计目的')
    body(doc, '《移动应用开发实践》是通信工程专业的一门重要实践课程，旨在培养学生掌握移动应用开发的完整流程，包括需求分析、系统设计、编码实现、测试调试以及团队协作等环节。本课程设计要求学生以小组协作的方式，综合运用前端开发、后端服务、数据库设计、AI集成等知识，完成一款功能完整、体验良好的移动应用产品。')
    body(doc, '通过本次课程设计，学生能够达到以下目标：第一，掌握鸿蒙（HarmonyOS）原生应用开发技术，包括ArkTS语言、ArkUI声明式UI框架、RelationalStore本地数据库等核心能力；第二，理解前后端分离架构的设计思想，掌握RESTful API的设计与调用方法；第三，了解AI大模型在移动应用中的集成方式，学习Dify低代码AI工作流平台的编排与调试技术；第四，培养团队协作能力和工程化开发素养，为后续毕业设计和职业发展奠定基础。')

    heading2(doc, '1.2  课程设计任务与要求')

    heading3(doc, '1.2.1  项目背景与意义')
    body(doc, '当代大学生群体普遍存在作息不规律、任务拖延、缺乏自我复盘习惯等问题。传统的待办事项应用功能单一，仅提供列表管理能力，缺少对用户行为的深层分析和个性化指导。与此同时，鸿蒙（HarmonyOS）操作系统正处于生态快速建设期，优质的鸿蒙原生应用在校园场景中存在大量需求缺口。')
    body(doc, '随着大语言模型（LLM）技术的成熟，AI赋能个人效率工具成为技术热点。然而，如何在移动应用中合理集成AI能力——既不增加用户使用负担，又能真正提升使用体验——是一个值得探索的工程命题。本项目"朝夕"——生活秩序管理助手，以鸿蒙原生开发为基础，整合Dify AI工作流与通义千问大模型，打造一款覆盖作息记录、待办管理、日记记录和游戏化成就激励的全方位个人生活管理应用。项目核心设计理念为"原生工具为核心，智能能力为辅助"，即所有基础功能100%离线可用，AI仅作为按需触发的增效插件。')

    heading3(doc, '1.2.2  课程设计任务')
    body(doc, '本课程设计任务为：以小组（3人）为单位，完成一款基于鸿蒙平台的移动应用的设计与开发。项目选题为"朝夕——生活秩序管理助手"，具体任务包括：')
    bullet(doc, '完成应用需求分析与系统设计，撰写需求规格说明和系统设计文档')
    bullet(doc, '使用ArkTS + ArkUI完成鸿蒙客户端开发，实现作息打卡、待办管理、日记记录和游戏化成就系统四大核心功能模块')
    bullet(doc, '设计并开发Python Flask中转服务器，解决Dify API Key安全问题和请求代理需求')
    bullet(doc, '在Dify平台上编排AI工作流，实现作息诊断、任务拆解、每日小结、月度复盘四种AI场景')
    bullet(doc, '完成前后端联调与系统测试，确保应用功能完整、运行稳定')
    bullet(doc, '按照学院模板格式撰写课程设计报告，系统总结开发过程与收获体会')

    heading3(doc, '1.2.3  课程设计要求')
    body(doc, '课程设计报告应按照学院统一模板格式撰写，包含完整的封面、目录、正文和参考文献。正文部分需涵盖需求分析、系统设计、系统实现、系统测试及总结体会等章节。因项目为小组协作完成，报告的前半部分（需求分析、总体设计）为小组共用内容，后半部分（详细实现和测试）应突出个人负责的工作内容。')
    body(doc, '本人（温志远）在小组中的主要职责为：第一，鸿蒙前端的整体UI设计与框架搭建，包括项目工程结构设计、页面路由导航框架、四个主页面及十余个业务组件的开发、数据模型层与服务层的架构设计，共计37个ArkTS源文件；第二，Python Flask中转服务器的完整设计与编码实现，包括核心转发逻辑、Web调试控制台（三栏布局单页应用）、自动转发与降级机制、请求与响应存档系统，共计456行Python代码。')

    # ============================================================
    # 第二章（Heading 1 自动换页）
    # ============================================================
    heading1(doc, '第二章  系统需求分析')

    heading2(doc, '2.1  用户需求分析')
    body(doc, '本应用的目标用户群体为在校大学生和职场新人。通过对目标用户的调研分析，总结出以下核心痛点：在作息方面，入睡时间不固定，缺乏规律作息的监督和反馈机制；在任务管理方面，待办事项容易堆积，缺乏优先级管理和目标拆解能力；在自我反思方面，没有定期记录和复盘的习惯，无法有效回溯个人成长轨迹；在行为指导方面，缺乏个性化的生活改进建议。')
    body(doc, '基于上述痛点分析，用户对应用的核心诉求包括：便捷的作息打卡功能，能够自动计算睡眠时长并判断是否熬夜或睡眠不足；灵活的待办任务管理，支持优先级分类和AI智能拆解大目标；简单的日记记录体验，支持心情标记和日历回顾；游戏化的成就激励系统，通过等级成长和徽章收集提升持续使用的动力；可选的AI智能分析功能，在需要时提供作息诊断、任务拆解建议、每日小结和月度复盘。')

    heading2(doc, '2.2  功能需求分析')

    heading3(doc, '2.2.1  作息健康记录')
    body(doc, '作息模块提供每日睡眠数据的录入与管理功能。用户可记录入睡时间、起床时间、睡眠质量评分（1-5分）、夜间醒来次数和入睡难易程度。系统自动计算睡眠时长，并根据预设规则判定是否熬夜（23:30后入睡）和睡眠不足（少于7小时）。支持按日、周、月查看历史作息记录的趋势变化，帮助用户直观了解自身睡眠状况。')

    heading3(doc, '2.2.2  待办任务管理')
    body(doc, '待办模块支持任务的完整生命周期管理（创建、编辑、删除、状态切换）。每个任务可设置标题、描述、三级优先级（高/中/低，分别以红/黄/绿色标签标识）、截止日期和父任务关联（子任务拆分）。任务列表支持按日期筛选，并按"未完成"和"已完成"分组展示。用户可对模糊的大目标使用"AI拆解"功能，系统自动生成结构化的子任务列表并支持一键导入。')

    heading3(doc, '2.2.3  日记记录')
    body(doc, '日记模块提供简洁的日记编辑功能，用户可为每篇日记选择心情标签（开心、平静、一般、低落），并以对应的emoji图标和颜色在日历视图中标记。日历视图按月份渲染，用户可一目了然地查看当月的情绪分布。日记详情页采用双区域布局：上方展示用户原创日记，下方展示AI生成的每日小结，实现"用户表达"与"AI总结"的对比阅读。')

    heading3(doc, '2.2.4  游戏化成就激励系统')
    body(doc, '成就系统通过即时正反馈机制帮助用户坚持良好习惯，包含三个激励维度：每日成就点（作息打卡+1、全部待办完成+1、撰写日记+1，每日上限3点）；六级等级体系（初入朝夕0点→秩序新手10点→节奏学徒30点→自律达人60点→生活大师100点→朝夕领航者150点）；八枚荣誉徽章（早起鸟·初阶/达人、好睡眠·初阶、任务粉碎机、全勤·初阶、日记达人/大师、完美一日），覆盖作息、待办、日记和综合四类。每次成就变动后，系统自动扫描徽章条件，满足条件时即时解锁，无需用户手动领取。')

    heading3(doc, '2.2.5  AI增效功能')
    body(doc, 'AI功能遵循"按需触发、不打扰"原则，所有AI能力仅在用户主动点击时运行。系统提供四种AI场景：作息诊断（分析近7天睡眠数据，输出健康评分、问题识别和分阶段改善方案）；任务拆解（将用户输入的模糊大目标拆解为可执行的子任务，含优先级、建议截止日期和实施建议）；每日小结（整合当日作息、任务完成情况和日记内容，生成结构化的当日总结）；月度复盘（全月四维度综合分析，提炼关键词并给出下月行动计划）。')

    heading2(doc, '2.3  非功能性需求')
    body(doc, '在性能方面，核心功能（作息打卡、待办操作、日记编辑）的操作响应时间应小于500毫秒，本地数据库查询时间小于200毫秒。在可靠性方面，核心功能必须100%离线可用，AI功能不可用时自动降级为Mock响应，不阻塞用户使用。在安全性方面，用户所有数据存储于设备本地RelationalStore，AI请求仅上传必要的统计摘要数据（不含日记原文全文或设备标识）。Dify API Key存储于中转服务器而非客户端，避免反编译泄露风险。在兼容性方面，应用需兼容HarmonyOS API 12及以上版本设备。')

    # ============================================================
    # 第三章
    # ============================================================
    heading1(doc, '第三章  系统设计')

    heading2(doc, '3.1  系统总体架构')
    body(doc, '系统采用三层架构设计，分别为客户端层（HarmonyOS App）、服务层（Python Flask中转站 + Dify AI平台）和数据层（本地RelationalStore + 云端知识库）。客户端层负责用户交互界面和本地数据管理，通过HTTP协议与中转服务器通信。中转服务器接收App的AI请求，转发至Dify工作流引擎处理后将结果返回给客户端。Dify平台运行单条工作流，通过场景参数路由到四条独立处理链路（每条链路包含数据预处理、LLM对话生成、知识库检索等节点），最终返回结构化JSON结果。')

    table_caption(doc, '表3-1  系统三层架构说明')
    create_table(doc,
        ['层级', '技术组件', '主要职责'],
        [
            ['客户端层', 'ArkTS + ArkUI', '用户交互界面、本地数据管理、AI请求发起'],
            ['服务层', 'Python Flask', 'API Key安全管理、请求转发代理、调试控制台'],
            ['AI层', 'Dify + 通义千问qwen3-max', '工作流编排、LLM推理、知识库检索（RAG）'],
            ['数据层', 'RelationalStore + Preferences', '5张业务表存储、键值对配置存储'],
        ]
    )

    heading2(doc, '3.2  技术选型')
    body(doc, '系统技术选型充分考虑了课程要求、生态成熟度和工程可行性。前端采用ArkTS语言配合ArkUI声明式框架，这是鸿蒙官方推荐的开发范式，能够充分发挥鸿蒙生态的原生能力。构建工具使用Hvigor（鸿蒙原生构建系统）。本地数据库采用RelationalStore存储5张核心业务表，轻量配置使用Preferences键值对存储。')
    body(doc, 'AI后端选择Dify低代码AI工作流平台，支持可视化工作流编排、知识库检索（RAG）和多模型接入，通过单工作流多分支设计将4种AI场景收敛到同一工作流，极大降低维护成本。AI模型选用通义千问qwen3-max。中转服务器采用Python Flask框架，轻量灵活，适合快速搭建API代理层。')

    table_caption(doc, '表3-2  技术选型汇总')
    create_table(doc,
        ['层级', '技术', '选型理由'],
        [
            ['客户端框架', 'ArkTS (HarmonyOS API 12+)', '课程要求，鸿蒙原生能力'],
            ['开发工具', 'DevEco Studio 5.0+', '鸿蒙官方IDE'],
            ['构建工具', 'Hvigor', '鸿蒙原生构建系统'],
            ['本地数据库', 'RelationalStore', '鸿蒙官方关系型数据库，类SQLite接口'],
            ['轻量存储', 'Preferences', '键值对存储，用于配置和成就总览'],
            ['AI平台', 'Dify工作流', '低代码编排，单工作流多场景路由'],
            ['AI模型', '通义千问 qwen3-max', '中文能力优秀，百炼平台'],
            ['中转服务器', 'Python Flask', '轻量灵活，解决API Key安全问题'],
            ['知识增强', 'RAG知识库检索', '睡眠健康知识库增强AI回答专业性'],
        ]
    )

    heading2(doc, '3.3  数据库设计')
    body(doc, '数据库采用鸿蒙RelationalStore，设计5张核心业务表，遵循第三范式。每张表均包含id自增主键和created_at/updated_at时间戳字段。5张表分别为：作息记录表（sleep_records）、待办任务表（tasks）、日记表（diaries）、成就记录表（achievement_records）和徽章表（badges）。')

    for title, fields in [
        ('表3-3  作息记录表（sleep_records）',
         [['id', 'INTEGER (PK)', '自增主键'], ['record_date', 'TEXT', '记录日期（YYYY-MM-DD）'],
          ['sleep_time', 'TEXT', '入睡时间（HH:mm）'], ['wake_time', 'TEXT', '起床时间（HH:mm）'],
          ['sleep_duration', 'REAL', '睡眠时长（小时）'], ['sleep_quality', 'INTEGER', '睡眠质量评分（1-5）'],
          ['wake_times', 'INTEGER', '夜间醒来次数'], ['sleep_difficulty', 'REAL', '入睡难易程度'],
          ['is_overnight', 'INTEGER', '是否熬夜（0/1）'], ['is_insufficient', 'INTEGER', '是否睡眠不足（0/1，<7h）'],
          ['created_at', 'TEXT', '创建时间戳'], ['updated_at', 'TEXT', '更新时间戳']]),
        ('表3-4  待办任务表（tasks）',
         [['id', 'INTEGER (PK)', '自增主键'], ['title', 'TEXT', '任务标题'],
          ['description', 'TEXT', '任务描述'], ['priority', 'INTEGER', '优先级（0=低 1=中 2=高）'],
          ['due_date', 'TEXT', '截止日期'], ['status', 'INTEGER', '状态（0=未完成 1=已完成）'],
          ['parent_task_id', 'INTEGER', '父任务ID'], ['created_at', 'TEXT', '创建时间戳'],
          ['updated_at', 'TEXT', '更新时间戳']]),
        ('表3-5  日记表（diaries）',
         [['id', 'INTEGER (PK)', '自增主键'], ['record_date', 'TEXT', '记录日期'],
          ['content', 'TEXT', '日记正文'], ['mood_label', 'TEXT', '心情标签（开心/平静/一般/低落）'],
          ['mood_emoji', 'TEXT', '心情emoji图标'], ['ai_text', 'TEXT', 'AI每日小结文本'],
          ['created_at', 'TEXT', '创建时间戳'], ['updated_at', 'TEXT', '更新时间戳']]),
    ]:
        table_caption(doc, title)
        create_table(doc, ['字段名', '类型', '说明'], fields)

    heading2(doc, '3.4  Dify AI工作流设计')
    body(doc, 'AI工作流采用"单工作流+场景参数驱动"的核心架构设计，是本项目的创新点之一。App端向中转站传入三个核心参数：scene_type（场景类型枚举）、user_info（用户基础信息：年龄、性别、爱好、习惯）和structured_data（各场景的结构化业务数据，JSON格式字符串）。中转站将请求转发至Dify后，工作流内部通过IF-ELSE条件分支节点根据scene_type的值路由到4条独立处理链路。')
    body(doc, '每条链路的处理流程为：Python Code节点（数据预处理与统计计算）→ LLM节点（通义千问qwen3-max进行推理和生成）→ 知识库检索节点（RAG，从睡眠健康知识库获取专业参考信息以增强回答质量）→ LLM节点（结合知识库检索结果生成最终回答）→ Python Code节点（格式标准化，确保输出JSON格式稳定可靠）。这种设计对外仅暴露1个API endpoint，4种场景共享同一套知识库和Prompt工程，新增场景只需在Dify工作流中增加一条分支链路，无需修改任何客户端代码。')

    heading2(doc, '3.5  安全架构设计')
    body(doc, '安全设计贯穿系统架构的各个层面。第一，Dify API Key不存储在客户端代码中，而是仅存在于中转服务器的配置变量中，客户端通过中转服务器间接调用Dify，从根本上消除了反编译获取API Key的风险。第二，所有用户数据完全存储于鸿蒙设备本地的RelationalStore，不上传至任何云端服务。第三，AI请求仅携带必要的统计摘要数据（如作息记录的日均值、任务完成率等），不包含日记原文全文或任何设备标识信息。第四，中转站设计了120秒超时自动降级机制，当Dify服务不可用或响应超时时，自动使用内置Mock响应模板返回给App，确保App端永远不会因AI服务异常而卡死。')

    # ============================================================
    # 第四章 — 个人负责重点
    # ============================================================
    heading1(doc, '第四章  系统实现')
    body(doc, '本章详细介绍系统的实现过程。其中，4.2节（鸿蒙前端UI设计与框架搭建）和4.3节（Python Flask中转服务器设计与实现）为本人（温志远）在小组中独立负责完成的核心工作内容。')

    heading2(doc, '4.1  开发环境搭建')
    body(doc, '鸿蒙客户端开发环境：DevEco Studio 5.0+ IDE，运行于macOS系统，目标SDK为HarmonyOS API 24+，构建工具为Hvigor。中转服务器开发环境：Python 3.x配合Flask框架和requests库，开发工具为VS Code。Dify工作流编排在Dify平台Web控制台中完成，通过API Key连接中转服务器。版本管理使用Git。')

    # 4.2 前端
    heading2(doc, '4.2  鸿蒙前端UI设计与框架搭建（个人负责）')
    body(doc, '鸿蒙前端的整体UI设计与框架搭建是本人负责的核心工作之一。前端工程共包含37个ArkTS源文件，采用分层架构设计，代码组织规范清晰。')

    table_caption(doc, '表4-1  前端代码分层统计')
    create_table(doc,
        ['层级', '目录', '文件数', '职责说明'],
        [
            ['页面层', 'pages/', '5', '页面UI布局与交互逻辑'],
            ['组件层', 'components/', '13', '可复用UI组件（base通用 + business业务）'],
            ['服务层', 'services/', '9', 'AI调用服务 + 8个数据源CRUD封装'],
            ['模型层', 'models/', '6', '数据模型定义与方法'],
            ['公共层', 'common/', '4', '常量定义、类型枚举、工具函数、成就管理器'],
        ]
    )

    heading3(doc, '4.2.1  项目工程结构')
    body(doc, '前端项目采用标准鸿蒙工程结构，主要代码位于entry/src/main/ets/目录下。pages/目录包含Index.ets（主入口+底部Tab容器）、HomePage.ets（首页）、TaskPage.ets（待办页）、DiaryPage.ets（日记页）、ProfilePage.ets（我的页）和SleepHistoryPage.ets（作息历史页）共6个页面文件。components/目录分为base子目录（BottomTab.ets底部导航、ModalSheet.ets通用弹窗）和business子目录（SleepRecordForm.ets作息表单、TaskItem.ets/TaskForm.ets/TaskActionsSheet.ets任务三件套、DiaryEditor.ets日记编辑器、CalendarView.ets日历视图、AchievementCard.ets成就卡片、AIResultPanel.ets AI结果面板、OnboardingQuestionnaire.ets引导问卷）共11个组件文件。')
    body(doc, 'services/目录按职责进一步分为storage子目录（DatabaseHelper.ets数据库初始化与表结构管理）、data子目录（SleepDataSource.ets等8个数据源，封装CRUD和高级查询）和api子目录（AIService.ets AI工作流通用调用）。这种分层架构使得代码职责清晰、依赖方向可控，在开发后期仍能保持良好的可维护性。')

    heading3(doc, '4.2.2  页面导航与路由框架')
    body(doc, '应用采用底部四Tab导航架构。入口页面Index.ets作为容器组件，使用@State状态变量currentPage控制四个子页面（home/task/diary/profile）的切换。底部导航组件BottomTab.ets封装四个Tab的图标、文字标签和点击切换逻辑，导航栏高度为62vp，设有选中态高亮样式。页面间的二级路由跳转（如从首页跳转到作息历史页SleepHistoryPage）使用router.pushUrl()实现，路由参数通过router.params传递。')
    body(doc, 'Index.ets的aboutToAppear()生命周期方法中执行了三项初始化任务：检查首次启动引导问卷状态（通过Preferences读取onboarding_completed标记）、初始化当日成就记录（调用AchievementManager.ensureTodayRecord()确保成就记录行存在）、检查是否需要自动生成每日总结（若当日尚未生成则触发AI每日小结流程）。')

    heading3(doc, '4.2.3  首页模块')
    body(doc, '首页（HomePage.ets）整合了个人问候、成就概览、作息记录、AI诊断、待办概览、日记概览和AI总结七个功能区域，信息架构由上至下依次排列。顶部的问候区域根据getGreeting()函数返回6个时段（夜深/早上/上午/中午/下午/晚上）的动态问候语，配合当前日期和模拟天气信息展示。')
    body(doc, '今日成就卡片以AchievementCard.ets组件渲染圆形进度条，直观展示距下一等级还需多少点数，同时列出当日三项成就目标（作息打卡、全部待办、日记）各自的完成状态。作息卡片展示最近一晚的入睡时间、起床时间和自动计算的睡眠时长，并以红色标签标注"熬夜"或"睡眠不足"状态，用户点击可弹出SleepRecordForm底部表单录入新数据。AI作息诊断卡片展示最近一次诊断结果（健康评分、识别问题、改善建议），内容支持长按复制。')
    body(doc, '待办和日记卡片分别提供当日的快速概览，点击可跳转至对应的功能页面。首页底部设置了AI功能双入口按钮（AI作息诊断和AI每日小结），用户点击后触发AIService对应方法，经中转服务器获取结果后通过AIResultPanel组件统一渲染展示。')

    heading3(doc, '4.2.4  待办任务模块')
    body(doc, '待办模块由TaskPage.ets页面和TaskItem.ets、TaskForm.ets、TaskActionsSheet.ets三个业务组件组成。TaskPage采用日期筛选栏+分段列表的布局方案：顶部日期选择器默认展示今天，用户可左右滑动切换日期；任务列表分为"未完成"和"已完成"两个分组，每个任务条目显示对应优先级的彩色标签（高=红色、中=黄色、低=绿色）。')
    body(doc, 'TaskItem.ets组件渲染单条任务的详细信息（标题、优先级色标、截止日期），右侧提供状态切换Toggle开关。点击条目弹出TaskActionsSheet底部操作面板，提供编辑、删除、添加子任务、AI拆解等操作的入口。TaskForm.ets是一个完整的表单组件，支持标题、描述输入，优先级下拉选择，截止日期选择器和父任务关联设置。AI拆解功能调用AIService.splitTask()方法，返回的子任务列表展示在弹窗中，用户可选择全部或部分导入到待办清单。')

    heading3(doc, '4.2.5  日记模块')
    body(doc, '日记模块由DiaryPage.ets页面、DiaryEditor.ets编辑器和CalendarView.ets日历视图组成，其中CalendarView是前端开发中技术难度较高的组件。CalendarView实现了月度日历的完整渲染逻辑：日期网格布局自动计算月初月末的空白填充、跨月日期的灰色淡化处理、心情颜色的精准映射（开心=绿色、平静=蓝色、一般=黄色、低落=灰色）、当天日期的高亮标识和月份快速切换功能。每个日期格根据当日是否有日记记录及对应的mood_label显示不同的背景色，用户点击任意日期可查看或编辑该天的日记。')
    body(doc, 'DiaryEditor.ets提供文本输入区域和心情标签选择器（四个emoji按钮：开心/平静/一般/低落）。DiaryPage采用独创的双区域对比布局：上方展示用户原创日记正文和心情标签，下方展示AI生成的每日小结（若已生成），实现用户表达与AI辅助总结的并排对比阅读。')

    heading3(doc, '4.2.6  个人中心与成就模块')
    body(doc, '个人中心页面（ProfilePage.ets）是成就系统和用户设置的展示中心。页面顶部展示用户头像、昵称、当前等级称号和成就点数，搭配圆形进度条直观显示距下一等级还需多少点数。中部为成就总览区域，展示累计打卡天数、完成任务总数、日记篇数、已解锁徽章数等统计指标。')
    body(doc, '页面底部的徽章墙以网格形式展示全部8枚徽章（4类：作息、待办、日记、综合），已解锁徽章以彩色图标和名称展示，未解锁徽章以灰色显示并标注解锁条件。AchievementCard.ets组件封装了单个成就卡片的展示逻辑。页面的AI月度复盘入口按钮触发AIService.monthlyReview()，生成涵盖睡眠健康、任务效率、生活状态和成就进展四维度的综合分析报告。')

    heading3(doc, '4.2.7  公共组件设计')
    body(doc, '为提升代码复用性和用户体验一致性，本项目设计和实现了以下公共组件：BottomTab.ets实现底部四Tab导航，支持选中态图标高亮和文字样式变化；ModalSheet.ets封装通用底部弹出面板，支持自定义标题、内容区域和操作按钮配置，被SleepRecordForm、TaskActionsSheet等多个业务场景复用；AIResultPanel.ets提供统一的AI结果渲染面板，支持树形结构展示、可编辑文本和复制到剪贴板功能，是四种AI场景的通用输出组件；SleepRecordForm.ets封装作息录入表单，用户在表单中修改时间后动态计算睡眠时长并实时判定熬夜和睡眠不足状态。')
    body(doc, '首次启动引导问卷组件OnboardingQuestionnaire.ets在用户首次启动应用时自动弹出，通过四步流程收集年龄、性别、爱好和习惯信息，结果存储到数据库中用于后续AI分析的个性化参考。引导完成后通过Preferences持久化标记onboarding_completed状态，后续启动不再重复显示。')

    heading3(doc, '4.2.8  数据模型与服务层')
    body(doc, '数据模型层（models/）定义了6个核心模型文件。SleepRecord.ets定义作息记录模型及CRUD输入/输出接口。Task.ets定义任务模型、优先级枚举和任务统计结构。Diary.ets定义日记模型及四种心情标签到emoji和颜色的映射配置。Achievement.ets定义成就记录模型、等级配置（6级所需点数阈值）和8枚徽章的完整预设数据（徽章key、名称、描述、图标、分类、解锁条件）。AIResult.ets定义AI四种场景的返回结果类型（SleepDiagnosisResult、TaskSplitResult、DailySummaryResult、MonthlyReviewResult）。Onboarding.ets定义引导问卷的用户信息模型。')
    body(doc, '服务层（services/）包含storage子层的DatabaseHelper.ets和data子层的8个数据源文件。DatabaseHelper使用单例模式管理全局RdbStore实例，首次调用getDatabase()时自动执行版本检查与表结构创建（5张表的DDL）。8个数据源文件分别封装对应业务表的完整CRUD操作和高级查询（按日期范围查询、聚合统计、连续天数计算等），所有数据库操作均返回Promise，确保UI线程不被阻塞。')
    body(doc, 'API调用服务AIService.ets封装了与中转服务器的完整通信逻辑。核心方法callWorkflow()接收sceneType、data和userInfo三个参数，自动组装符合Dify API格式的请求体，通过@kit.NetworkKit的http模块发送POST请求（30秒连接超时、60秒读取超时）。响应经过3层JSON解析（HTTP响应→DifyWorkflowResponse→UnifiedAIResult→场景具体类型），每层均设有异常兜底，解析失败时返回明确的错误信息而非应用崩溃。4个场景方法（diagnoseSleep、splitTask、dailySummary、monthlyReview）均为对callWorkflow的封装调用。')

    # 4.3 中转站
    heading2(doc, '4.3  Python Flask中转服务器设计与实现（个人负责）')
    body(doc, 'Python Flask中转服务器是本项目的另一项个人核心工作，也是本系统安全架构的关键组件。该服务器解决了Dify API Key不能存储在客户端代码中的安全难题，同时承担请求存档、响应存档、Web调试控制台、自动转发降级等多重职责。整个服务器共计456行Python代码，包含完整的Flask路由、Dify转发引擎和单页Web控制台。')

    heading3(doc, '4.3.1  中转站架构设计')
    body(doc, '中转站的核心架构围绕"请求→转发→响应"三段式流水线设计。当鸿蒙App发起AI请求时，数据流如下：首先，请求到达中转站的/api/workflows/run路由，路由解析请求体并保存至requests/目录；接着，根据配置（自动转发开关状态）决定是否将请求转发至Dify API；最后，等待最终响应内容确定（Dify真实响应、控制台手动编辑内容或超时降级Mock响应）后返回给App。')
    body(doc, '中转站的核心组件包括：Flask HTTP服务器（监听4567端口，支持多线程处理）、待响应请求队列（pending_requests字典，配合threading.Lock保证线程安全，threading.Event实现阻塞等待-唤醒同步）、Dify转发引擎（do_forward_to_dify函数，可被路由手动触发或后台自动转发线程调用）、Web调试控制台（内联单页HTML/CSS/JS，三栏布局通过REST API与后端通信）。')

    heading3(doc, '4.3.2  核心路由与请求处理')
    body(doc, '/api/workflows/run（POST）是中转站的核心路由。处理流程分为五步：第一步，通过全局计数器生成递增请求序号（格式0001、0002等），用于后续追溯。第二步，解析请求体JSON，提取scene_type场景类型和structured_data结构化数据。第三步，调用save_request_file()将请求元信息、原始JSON和解析后的数据保存为带时间戳的JSON文件（命名格式：序号_时间戳_场景类型.json）。第四步，将请求加入pending_requests队列并创建threading.Event对象用于阻塞等待。第五步，根据自动转发开关状态决定行为——若已开启则在后台线程中自动调用do_forward_to_dify()，若关闭则在终端打印提示等待用户在控制台中手动操作。最后调用event.wait(timeout=120)阻塞等待响应就绪，超时后自动使用默认Mock响应返回给App。')

    heading3(doc, '4.3.3  Dify转发引擎')
    body(doc, 'do_forward_to_dify(req_id)函数是中转站的核心引擎。该函数执行流程：首先进行前置检查（验证ENABLE_DIFY_FORWARD开关和DIFY_API_KEY是否已配置），检查通过后从待响应队列中取出请求原始数据。然后使用Python requests库向Dify API发送HTTP POST请求，携带Bearer Token认证头和JSON请求体，超时设为90秒。转发完成后记录HTTP状态码和耗时（毫秒精度），调用save_response_file()将Dify原始响应保存为JSON文件至response/目录。同时更新队列中该请求的状态字段（dify_response、dify_status、dify_elapsed_ms）。若HTTP状态码为200，将Dify响应设为最终返回给App的响应文本。整个过程使用pending_lock互斥锁保证线程安全。')

    heading3(doc, '4.3.4  Web调试控制台')
    body(doc, 'Web调试控制台通过浏览器访问http://localhost:4567/即可打开，是开发调试的重要工具。控制台采用三栏响应式布局设计：左栏为实时更新的请求列表（展示序号、场景类型、数据大小和状态标签，每2秒自动轮询刷新）；中栏上半部分为请求数据查看区，支持三种视图模式切换（树形结构化数据、字段平铺列表、原始JSON代码），中栏下半部分为Dify响应面板；右栏为响应编辑器，内置JSON文本编辑区，用户可手动编辑响应内容并提供发送、快速发送、重置默认值和JSON格式校验四个操作按钮。')
    body(doc, 'Web控制台的完整HTML/CSS/JS代码（约570行）以内联方式通过Flask的render_template_string()嵌入在server_relay_updata.py文件中，实现单文件部署，无需额外静态资源。界面采用深色主题（深蓝灰背景色），与开发工具风格保持一致。前端异步操作通过fetch API与后端多个REST接口通信，实现完整的请求→转发→编辑→发送调试工作流。')

    heading3(doc, '4.3.5  自动转发与降级机制')
    body(doc, '中转站设计了双重自动化开关系统，通过控制台顶部的两个Toggle组件控制。第一个开关"自动转发Dify"：开启后新请求到达时立即在后台线程中自动转发至Dify；第二个开关"自动转发响应到App"：Dify返回响应后立即发送给等待中的App，整个AI请求-响应流程完全自动化。两个开关的状态通过/api/auto-forward和/api/auto-send接口读写，使用auto_lock线程锁保证并发环境下的状态一致性。')
    body(doc, '降级机制确保系统在任何异常情况下都不会阻塞用户。第一层降级：Dify API Key未配置或转发未启用时，build_default_response()函数根据场景类型生成内置Mock响应模板（覆盖4种AI场景的合理模拟数据）。第二层降级：请求在队列中等待超过120秒时，event.wait()超时返回，路由自动使用当前response_text（即默认Mock响应）返回给App。这套双重降级机制确保了核心功能在Dify不可用时仍能提供静态建议，用户体验不受影响。')

    heading3(doc, '4.3.6  请求与响应存档')
    body(doc, '中转站自动将所有AI请求和Dify响应存档为结构化JSON文件，分别保存在requests/和response/目录下。每个文件以"序号_时间戳_场景类型.json"格式命名（如0001_20260705_143022_sleep_diagnosis.json），便于按时间序追溯。请求存档包含元信息、原始请求体和解析后的structured_data；响应存档包含元信息、原始响应文本和解析后的响应JSON。这套存档机制使得每次AI交互均可被完整回顾和分析，极大提升了调试效率和问题定位能力。')

    # 4.4 AI联调
    heading2(doc, '4.4  AI功能集成与联调')
    body(doc, 'AI功能的前后端联调是系统实现的重要环节。鸿蒙客户端通过DIFY_API_URL常量（配置为http://10.0.2.2:4567/api/workflows/run，10.0.2.2是Android/鸿蒙模拟器中映射到宿主机localhost的特殊IP地址）连接中转服务器。AIService提供4个场景专用静态方法（diagnoseSleep、splitTask、dailySummary、monthlyReview），每个方法内部调用通用callWorkflow()并传入对应的SceneType枚举值和场景数据。')
    body(doc, '联调过程中，中转站Web控制台发挥了关键作用。开发者可实时查看App发送的请求内容是否完整正确、Dify返回的响应格式是否符合预期，必要时可在控制台中手动编辑响应文本后再发送给App。这种"中间人可干预"的架构设计显著降低了前后端开发的耦合度——当Dify工作流仍在调试优化时，App端的开发和测试完全不受影响，通过Mock响应即可独立推进。')

    # ============================================================
    # 第五章
    # ============================================================
    heading1(doc, '第五章  系统测试')

    heading2(doc, '5.1  功能测试')
    body(doc, '对各功能模块进行了系统性测试，覆盖正常流程和边界条件，主要测试结果如表5-1所示。作息打卡模块可正确录入睡眠数据，自动计算睡眠时长并判定熬夜/不足状态，历史记录查询和趋势展示正常。待办模块支持完整CRUD操作，优先级分类和分组展示正常，AI拆解功能可正确生成子任务并导入。日记模块支持文本编辑、心情标签选择和日历视图颜色标记，双区域展示正常。成就系统可正确计算每日成就点、自动检测徽章解锁条件、等级升级平滑过渡。')

    table_caption(doc, '表5-1  核心功能测试用例及结果')
    create_table(doc,
        ['测试模块', '测试用例', '预期结果', '结果'],
        [
            ['作息打卡', '录入入睡23:30、起床7:00', '判定未熬夜、睡眠充足', '通过'],
            ['作息打卡', '录入入睡01:00、起床6:00', '判定熬夜、睡眠不足', '通过'],
            ['待办管理', '创建高优任务并切换状态', '红色标签，状态正确切换', '通过'],
            ['待办管理', 'AI拆解"准备期末考试"', '返回3-5个子任务，可导入', '通过'],
            ['日记记录', '选择心情"开心"并保存', '日历视图中日期显示绿色', '通过'],
            ['成就系统', '完成三项日常打卡', '获得3点成就，进度条更新', '通过'],
            ['徽章检测', '连续7天8:00前起床', '自动解锁早起鸟·初阶徽章', '通过'],
            ['AI降级', '关闭Dify服务后触发AI', '120s超时，返回Mock响应', '通过'],
        ]
    )

    heading2(doc, '5.2  兼容性测试')
    body(doc, '在DevEco Studio模拟器（HarmonyOS API 12）上进行了完整兼容性测试。应用在手机竖屏模式下的所有页面布局显示正常，组件间距和字号适配良好。本地数据库RelationalStore读写操作运行稳定，未出现数据丢失或损坏情况。底部Tab切换动画流畅，页面间路由跳转无异常，应用内存占用正常。')

    heading2(doc, '5.3  AI集成测试')
    body(doc, 'AI功能测试覆盖了正常响应和异常降级两种场景。在正常场景下：中转站正确接收App请求（请求体格式验证通过），成功转发至Dify API（HTTP 200返回），Dify工作流按scene_type正确路由到对应处理链路，返回的结构化JSON经3层解析后成功展示在App界面。在异常降级场景下（模拟Dify服务不可用）：中转站正确识别异常状态，120秒超时后自动使用内置Mock响应模板返回给App，App端正常展示预设内容，用户操作完全不受影响。中转站Web控制台在所有测试场景下均正确展示请求详情、转发状态和响应数据。')

    # ============================================================
    # 第六章
    # ============================================================
    heading1(doc, '第六章  课程设计总结与体会')
    body(doc, '通过本次移动应用开发实践课程设计，我完成了"朝夕"——生活秩序管理助手项目中鸿蒙前端UI设计、框架搭建以及Python Flask中转服务器的全套设计与开发工作。回顾整个开发过程，我在多个技术维度上都获得了显著的成长。')
    body(doc, '在鸿蒙前端开发方面，我从零学习了ArkTS语言和ArkUI声明式UI框架。与传统的Android XML布局开发相比，ArkUI的声明式编程范式更加简洁直观，@State和@Prop等状态管理机制有效降低了UI与数据之间的同步复杂度。在搭建项目工程结构时，我遵循分层架构设计原则，将37个源文件按页面、组件、服务、模型、公共五个层级组织，每个层级职责明确、依赖方向可控。这种架构使得代码在开发后期仍能保持良好的可维护性和可扩展性。')
    body(doc, '日历视图组件（CalendarView.ets）的开发是前端工作中技术挑战最大的部分。组件需要处理月度日历的日期网格渲染逻辑——包括月初月末的空白填充、跨月日期的灰色淡化处理、心情颜色到日期格的精准映射和当天日期的高亮标识——这些细节对布局系统和条件渲染机制的理解要求较高。通过反复调试和优化，该组件最终实现了流畅的渲染效果和良好的交互体验。')
    body(doc, '在Python Flask中转服务器的开发中，我重点解决了两个关键问题。第一是API Key安全问题：通过引入中间代理层，将敏感的Dify API Key锁定在服务端内存中，客户端完全无感知。这使我深刻理解了移动应用安全架构设计的基本原则——任何存储在客户端的敏感信息都存在被反编译泄露的风险，密钥凭证必须由服务端持有。第二是系统可靠性问题：通过120秒超时降级机制和内置Mock响应模板，确保AI服务不可用时不影响用户使用核心功能。这种"优雅降级"的设计思想不仅适用于AI功能集成场景，也是构建稳健软件系统的通用工程原则。')
    body(doc, '中转站Web调试控制台的开发也带来了多方面的收获。我将HTML/CSS/JS代码全部以内联方式嵌入Python源文件中，实现了单文件部署——无需额外静态资源目录，一个Python脚本即可启动完整的调试面板。三栏布局、深色主题、实时轮询、树形JSON展示等功能虽然实现细节繁多，但最终呈现的调试体验非常高效，在前后端联调阶段发挥了不可替代的作用。双开关自动化（自动转发+自动发送）的设计则体现了我对"降低重复操作、提升开发效率"的思考。')
    body(doc, '在团队协作方面，本次课程设计采用3人小组分工模式。我负责的前端界面和中转服务器是连接用户和后端AI能力的桥梁，因此需要与负责Dify工作流编排和数据库设计的同学密切配合。我们约定好API接口的请求/响应JSON Schema格式，各自按接口规范独立开发，最终在中转站Web控制台中进行联调验证。这种"基于接口契约的并行开发"模式使得各模块可以互不阻塞地推进，显著缩短了整体开发周期。')
    body(doc, '在开发过程中也遇到了一些问题和挑战。例如，初期前端首页加载速度偏慢，排查后发现是因为数据库查询未做分页处理，一次性加载了全部历史数据，改为按需分页查询后性能显著改善。又如，中转站Web控制台最初未设计自动转发开关，每次新请求到达都需要手动点击转发按钮，在连续测试多个场景时操作繁琐，后续版本通过增加双开关自动化功能解决了这一效率瓶颈。')
    body(doc, '综上所述，本次课程设计让我系统性地实践了移动应用开发的全流程——从需求分析、架构设计到编码实现、测试调试。特别是在鸿蒙原生开发和AI大模型应用集成这两个前沿技术方向上，积累了宝贵的实战经验。这些收获不仅帮助我顺利完成了课程设计任务，也为未来的毕业设计和职业发展奠定了坚实的基础。')

    # ============================================================
    # 参考文献
    # ============================================================
    heading1(doc, '参考文献')
    refs = [
        '[1] 华为技术有限公司. HarmonyOS应用开发文档[EB/OL]. https://developer.huawei.com/consumer/cn/doc/harmonyos-guides, 2024.',
        '[2] 华为技术有限公司. ArkTS语言规范[EB/OL]. https://developer.huawei.com/consumer/cn/doc/harmonyos-guides/arkts-spec, 2024.',
        '[3] Dify.AI. Dify平台使用文档[EB/OL]. https://docs.dify.ai/zh-hans, 2024.',
        '[4] 阿里云计算有限公司. 通义千问模型使用文档[EB/OL]. https://help.aliyun.com/document_detail/tongyi.html, 2024.',
        '[5] Grinberg M. Flask Web开发:基于Python的Web应用开发实战[M]. 安道, 译. 北京: 人民邮电出版社, 2018.',
        '[6] Brown E. Web开发技术:HTML、CSS与JavaScript实战[M]. 北京: 机械工业出版社, 2022.',
        '[7] 张海藩, 吕云翔. 软件工程(第6版)[M]. 北京: 人民邮电出版社, 2020.',
        '[8] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. Irvine: University of California, 2000.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Pt(0)
        _set_exact_line_spacing(p, 22)
        _add_run(p, ref, size=Pt(12))

    # ============================================================
    doc.save(OUTPUT_FILE)
    print(f'✅ 课程设计报告已生成: {OUTPUT_FILE}')
    print(f'   ⚠️  学号占位符: {STUDENT_ID}')
    print(f'   📋  打开 Word 后 → 右键目录 → 更新域 → 更新整个目录')
    print(f'   📋  Heading 1/2/3 样式已设置，可在导航窗格中查看文档结构')

if __name__ == '__main__':
    generate()
