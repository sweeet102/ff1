#!/usr/bin/env python3
"""Fix course report page numbering - correct section break placement"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

filepath = '/Users/wenzhiyuan/Desktop/ff1/output/光纤通信课程报告_FWM_jiangsu-formatted.docx'
doc = Document(filepath)
body = doc.element.body
all_paras = list(body.iterchildren(qn('w:p')))

# Find first real body chapter (Heading 1 that's NOT 摘要/Abstract/目录/参考文献/致谢)
front_matter = {'摘要', 'Abstract', '目录', '参考文献', '致谢', '目  录', '摘  要', '致  谢'}
body_start = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'Heading 1' and text not in front_matter and i > 15:
        # Check it's a real chapter title (starts with 第 or number)
        if '第' in text or text[0].isdigit():
            body_start = i
            print('Body starts at para ' + str(i+1) + ': ' + text[:60])
            break

if body_start is None:
    print('ERROR: Could not find body start')
    exit(1)

# Remove any existing section breaks
all_paras = list(body.iterchildren(qn('w:p')))
for p in all_paras:
    pPr = p.find(qn('w:pPr'))
    if pPr is not None:
        old_sp = pPr.find(qn('w:sectPr'))
        if old_sp is not None:
            pPr.remove(old_sp)
            print('Removed existing section break from para')

# Insert section break just before the body chapter
prev_para = all_paras[body_start - 1]
pPr = prev_para.find(qn('w:pPr'))
if pPr is None:
    pPr = OxmlElement('w:pPr')
    prev_para.insert(0, pPr)

new_sectPr = OxmlElement('w:sectPr')

pgSz = OxmlElement('w:pgSz')
pgSz.set(qn('w:w'), '11906')
pgSz.set(qn('w:h'), '16838')
new_sectPr.append(pgSz)

pgMar = OxmlElement('w:pgMar')
pgMar.set(qn('w:top'), '1440')
pgMar.set(qn('w:bottom'), '1440')
pgMar.set(qn('w:left'), '1260')
pgMar.set(qn('w:right'), '1106')
pgMar.set(qn('w:header'), '851')
pgMar.set(qn('w:footer'), '992')
new_sectPr.append(pgMar)

# Section 1 (front matter): Roman numerals
pgNumType1 = OxmlElement('w:pgNumType')
pgNumType1.set(qn('w:fmt'), 'upperRoman')
pgNumType1.set(qn('w:start'), '1')
new_sectPr.append(pgNumType1)

pPr.append(new_sectPr)
print('Section break inserted before paragraph ' + str(body_start))

# Fix body section (last sectPr in body): Arabic numerals
body_sectPr = body.find(qn('w:sectPr'))
if body_sectPr is not None:
    # Remove any existing pgNumType
    for old_pn in body_sectPr.findall(qn('w:pgNumType')):
        body_sectPr.remove(old_pn)
    pgNumType2 = OxmlElement('w:pgNumType')
    pgNumType2.set(qn('w:fmt'), 'decimal')
    pgNumType2.set(qn('w:start'), '1')
    body_sectPr.insert(0, pgNumType2)
    print('Body section pgNumType set to decimal')

doc.save(filepath)

# ===== Re-open and set footers =====
doc2 = Document(filepath)
print('Sections: ' + str(len(doc2.sections)))

for i, sec in enumerate(doc2.sections):
    print('Section ' + str(i+1) + ':')
    # Footer
    footer = sec.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run._element.append(instr)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_end)
    print('  Footer: PAGE field added')

    # Header
    header = sec.header
    header.is_linked_to_previous = False
    for ph in header.paragraphs:
        ph.clear()
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rh = hp.add_run('江苏大学本科毕业设计（论文）')
    rh.font.size = Pt(9)
    rh.font.name = '宋体'
    print('  Header: school name added')

    # Different first page for section 1 (cover no number)
    if i == 0:
        sec.different_first_page_header_footer = True

doc2.save(filepath)
print('\nDone!')
print('Section 1 (封面+摘要+Abstract+目录): Roman numerals I, II, III...')
print('Section 2 (第一章 ~ 致谢): Arabic numerals 1, 2, 3...')
