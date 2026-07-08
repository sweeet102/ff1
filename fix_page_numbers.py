#!/usr/bin/env python3
"""Fix page numbering for both thesis files"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def find_first_chapter(doc):
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if p.style.name == 'Heading 1':
            if i > 15:
                return i
        if text and i > 15:
            if text.startswith('1 ') or (text.startswith('第') and '章' in text):
                return i
    return None

def insert_section_break(prev_para_elem):
    pPr = prev_para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        prev_para_elem.insert(0, pPr)

    new_sectPr = OxmlElement('w:sectPr')

    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    new_sectPr.append(pgSz)

    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1440')
    pgMar.set(qn('w:bottom'), '1440')
    pgMar.set(qn('w:left'), '1260')
    pgMar.set(qn('w:right'), '1106')
    pgMar.set(qn('w:header'), '851')
    pgMar.set(qn('w:footer'), '992')
    new_sectPr.append(pgMar)

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'upperRoman')
    pgNumType.set(qn('w:start'), '1')
    new_sectPr.append(pgNumType)

    pPr.append(new_sectPr)
    return new_sectPr

def set_body_pgNum(body_sectPr):
    pn = body_sectPr.find(qn('w:pgNumType'))
    if pn is None:
        pn = OxmlElement('w:pgNumType')
        body_sectPr.insert(0, pn)
    pn.set(qn('w:fmt'), 'decimal')
    pn.set(qn('w:start'), '1')

def add_page_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run._element.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_end)

def add_header_right(section, text):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = '宋体'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)

def fix_file(filepath, label):
    sep = '=' * 60
    print(sep)
    print('修复: ' + label)
    print('文件: ' + filepath)

    doc = Document(filepath)
    body = doc.element.body
    all_paras = list(body.iterchildren(qn('w:p')))

    chap_idx = find_first_chapter(doc)
    if chap_idx is None:
        print('  WARNING: 找不到第一章')
        return

    chap_text = doc.paragraphs[chap_idx].text[:60]
    print('  第一章位于段落 ' + str(chap_idx + 1) + ': ' + chap_text)

    # Check existing section breaks
    existing_breaks = []
    for idx, p in enumerate(all_paras):
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            sp = pPr.find(qn('w:sectPr'))
            if sp is not None:
                existing_breaks.append(idx)

    print('  现有分节符: ' + str(len(existing_breaks)) + ' 个 ' + str(existing_breaks))

    if chap_idx - 1 not in existing_breaks:
        prev_para = all_paras[chap_idx - 1]
        insert_section_break(prev_para)
        print('  已在段落 ' + str(chap_idx) + ' 前插入分节符')
    else:
        prev_para = all_paras[chap_idx - 1]
        pPr = prev_para.find(qn('w:pPr'))
        sp = pPr.find(qn('w:sectPr'))
        pn = sp.find(qn('w:pgNumType'))
        if pn is None:
            pn = OxmlElement('w:pgNumType')
            sp.insert(0, pn)
        pn.set(qn('w:fmt'), 'upperRoman')
        pn.set(qn('w:start'), '1')
        print('  已修复分节符 pgNumType -> upperRoman')

    # Fix body section
    body_sectPr = body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        set_body_pgNum(body_sectPr)
        print('  正文 pgNumType -> decimal')

    # Re-read and fix via python-docx sections
    doc2 = Document(filepath)
    print('  Sections: ' + str(len(doc2.sections)))

    if len(doc2.sections) >= 2:
        sec1 = doc2.sections[0]
        sec1.different_first_page_header_footer = True
        add_page_footer(sec1)
        # Front matter header: left=chapter name + right=school name
        # For simplicity, just school name on right
        header1 = sec1.header
        header1.is_linked_to_previous = False
        for ph in header1.paragraphs:
            ph.clear()
        p1h = header1.paragraphs[0] if header1.paragraphs else header1.add_paragraph()
        p1h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1h = p1h.add_run('江苏大学本科毕业设计（论文）')
        r1h.font.size = Pt(9)
        r1h.font.name = '宋体'

        sec2 = doc2.sections[1]
        sec2.different_first_page_header_footer = False
        add_page_footer(sec2)
        add_header_right(sec2, '江苏大学本科毕业设计（论文）')

        print('  页脚页眉已设置')
    elif len(doc2.sections) == 1:
        sec1 = doc2.sections[0]
        sec1.different_first_page_header_footer = True
        add_page_footer(sec1)
        add_header_right(sec1, '江苏大学本科毕业设计（论文）')
        print('  单section页脚已设置')

    doc2.save(filepath)
    print('  已保存')

# Fix both files
fix_file(
    '/Users/wenzhiyuan/Desktop/3231303014-王奇坤_formatted.docx',
    '王奇坤 - FBG论文'
)

fix_file(
    '/Users/wenzhiyuan/Desktop/ff1/output/光纤通信课程报告_FWM_jiangsu-formatted.docx',
    '光纤通信课程报告 - FWM'
)

print('=' * 60)
print('全部修复完成')
print('前置部分（封面/摘要/目录）-> 罗马数字 I II III')
print('正文部分（第一章起）-> 阿拉伯数字 1 2 3')
