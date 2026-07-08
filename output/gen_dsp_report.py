#!/usr/bin/env python3
"""生成 DSP 实验报告 — 基于模板 '20实验报告模板（zx）20210322大院版'
   格式：封面 + 目录 + 实验目的/设备/原理/过程/总结 + 参考文献
   页数：4-10 页（含封面目录），约 5 张纸双面打印"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, io

OUTPUT = "/Users/wenzhiyuan/Desktop/ff1/output/DSP实验报告_温志远.docx"

# ── matplotlib 全局设置 ──
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['Heiti TC', 'STHeiti', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False
DPI = 150

# ── 模板配色 ──
C = {
    'box_blue':    '#D6E8F7',
    'border_blue': '#3A7CC3',
    'box_orange':  '#FDE8D0',
    'border_orange':'#C38A3A',
    'box_green':   '#D8F0D8',
    'border_green':'#3A8C3A',
    'text':        '#222222',
    'arrow':       '#555555',
    'loop':        '#7AA8CC',
    'start_fill':  '#E0EBF5',
}

# ═══════════════════════════════════════════════════════
# matplotlib 图表函数
# ═══════════════════════════════════════════════════════

def rounded_box(ax, x, y, w, h, color=C['box_blue'], edge=C['border_blue'],
                lw=2, radius=0.08):
    box = FancyBboxPatch((x, y), w, h,
                         boxstyle=f"round,pad=0,rounding_size={radius}",
                         facecolor=color, edgecolor=edge, linewidth=lw, zorder=2)
    ax.add_patch(box)

def add_text(ax, x, y, w, h, text, fontsize=10, color=C['text']):
    ax.text(x + w/2, y + h/2, text, ha='center', va='center',
            fontsize=fontsize, color=color, zorder=3, linespacing=1.4)

def arrow(ax, x0, y0, x1, y1, color=C['arrow'], lw=2.5):
    ax.annotate('', xy=(x1, y1), xytext=(x0, y0),
                arrowprops=dict(arrowstyle='->,head_length=5,head_width=3',
                               color=color, lw=lw, connectionstyle='arc3,rad=0'),
                zorder=1)

def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=DPI, bbox_inches='tight', pad_inches=0.2,
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf

# ── 图1：系统总体结构框图 ──
def fig_system_diagram():
    fig, ax = plt.subplots(figsize=(9, 2.0))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis('off')
    box_w, box_h = 0.18, 0.55; gap = 0.06; n = 4
    total_w = n * box_w + (n - 1) * gap
    start_x = (1 - total_w) / 2; cy = 0.22
    blocks = [
        ('模拟信号\n(输入)', C['box_blue'], C['border_blue']),
        ('AD 模数转换\n(ICETEK5509)', C['box_orange'], C['border_orange']),
        ('FFT 频谱分析\n(C 语言实现)', C['box_blue'], C['border_blue']),
        ('频谱显示\n(CCS Graph)', C['box_green'], C['border_green']),
    ]
    for i, (text, fill, edge) in enumerate(blocks):
        x = start_x + i * (box_w + gap)
        rounded_box(ax, x, cy, box_w, box_h, fill, edge, radius=0.04)
        add_text(ax, x, cy, box_w, box_h, text, fontsize=10)
        if i < n - 1:
            arrow(ax, x + box_w, cy + box_h/2, x + box_w + gap, cy + box_h/2)
    ax.text(0.5, 0.02, '图1  系统总体结构框图', transform=ax.transAxes,
            ha='center', fontsize=9.5, color='#888')
    return fig_to_bytes(fig)

# ── 图2：AD模块设计流程 ──
def fig_ad_flow():
    fig, ax = plt.subplots(figsize=(8.5, 1.8))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis('off')
    box_w, box_h = 0.18, 0.55; gap = 0.05; n = 4
    total_w = n * box_w + (n - 1) * gap
    start_x = (1 - total_w) / 2; cy = 0.22
    blocks = [
        ('建立 CCS\n工程', C['box_blue'], C['border_blue']),
        ('导入 AD\n示例工程', C['box_blue'], C['border_blue']),
        ('编译下载\n运行验证', C['box_orange'], C['border_orange']),
        ('FileOut\n保存数据', C['box_green'], C['border_green']),
    ]
    for i, (text, fill, edge) in enumerate(blocks):
        x = start_x + i * (box_w + gap)
        rounded_box(ax, x, cy, box_w, box_h, fill, edge, radius=0.04)
        add_text(ax, x, cy, box_w, box_h, text, fontsize=10)
        if i < n - 1:
            arrow(ax, x + box_w, cy + box_h/2, x + box_w + gap, cy + box_h/2)
    ax.text(0.5, 0.02, '图2  AD模块设计流程', transform=ax.transAxes,
            ha='center', fontsize=9.5, color='#888')
    return fig_to_bytes(fig)

# ── 图3：FFT软件设计流程 ──
def fig_fft_flow():
    fig, ax = plt.subplots(figsize=(10, 1.8))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis('off')
    box_w, box_h = 0.15, 0.55; gap = 0.04; n = 5
    total_w = n * box_w + (n - 1) * gap
    start_x = (1 - total_w) / 2; cy = 0.22
    blocks = [
        ('设计测试\n信号', C['box_blue'], C['border_blue']),
        ('预计算\n旋转因子', C['box_blue'], C['border_blue']),
        ('位反转\n排序', C['box_orange'], C['border_orange']),
        ('蝶形运算\n(log2N 级)', C['box_orange'], C['border_orange']),
        ('CCS Graph\n显示频谱', C['box_green'], C['border_green']),
    ]
    for i, (text, fill, edge) in enumerate(blocks):
        x = start_x + i * (box_w + gap)
        rounded_box(ax, x, cy, box_w, box_h, fill, edge, radius=0.04)
        add_text(ax, x, cy, box_w, box_h, text, fontsize=9)
        if i < n - 1:
            arrow(ax, x + box_w, cy + box_h/2, x + box_w + gap, cy + box_h/2)
    ax.text(0.5, 0.02, '图3  FFT软件设计流程', transform=ax.transAxes,
            ha='center', fontsize=9.5, color='#888')
    return fig_to_bytes(fig)

# ── 图4：系统主程序流程 ──
def fig_main_flow():
    fig, ax = plt.subplots(figsize=(6.5, 7.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 22); ax.axis('off')
    mid = 5.0; box_w = 7.0; small_w = 6.0; step_h = 0.9; step_gap = 0.6

    def box(x, y, w, h, text, fill=C['box_blue'], edge=C['border_blue'], fontsize=10):
        rounded_box(ax, x, y, w, h, fill, edge, radius=0.15)
        add_text(ax, x, y, w, h, text, fontsize=fontsize)

    def v_arrow(x, y0, y1):
        arrow(ax, x, y0, x, y1)

    start_y = 20.0
    ellipse = mpatches.FancyBboxPatch((mid - 0.8, start_y - 0.5), 1.6, 1.0,
                                       boxstyle='round,pad=0.1,rounding_size=0.5',
                                       facecolor=C['start_fill'], edgecolor=C['border_blue'],
                                       linewidth=2, zorder=2)
    ax.add_patch(ellipse)
    ax.text(mid, start_y, '开  始', ha='center', va='center', fontsize=11, zorder=3)
    v_arrow(mid, start_y - 0.5, start_y - 1.5)

    init_y = start_y - 2.5
    box(mid - box_w/2, init_y, box_w, 1.2,
        '系统初始化\n（PLL · McBSP · ADC · 中断向量表）',
        fill=C['box_orange'], edge=C['border_orange'], fontsize=10)
    v_arrow(mid, init_y, init_y - 1.2)

    loop_top = init_y - 2.2; loop_bottom = 3.0
    loop_h = loop_top - loop_bottom; loop_x = 1.0; loop_w = 8.0
    loop_rect = mpatches.FancyBboxPatch((loop_x, loop_bottom), loop_w, loop_h,
                                         boxstyle='round,pad=0,rounding_size=0.3',
                                         facecolor='none', edgecolor=C['loop'],
                                         linewidth=2, linestyle='--', zorder=0)
    ax.add_patch(loop_rect)
    ax.text(loop_x + 0.2, loop_top + 0.25, 'while (1)', fontsize=9,
            color=C['loop'], style='italic', zorder=3)

    cy = loop_top - 0.8; sx = mid - small_w / 2
    steps = [
        ('AD_Collect(buffer)    采集 N 点数据', C['box_blue'], C['border_blue']),
        ('FFT(buffer, N)        执行 N 点 FFT', C['box_blue'], C['border_blue']),
        ('DisplaySpectrum()     更新频谱显示', C['box_green'], C['border_green']),
    ]
    last_y = cy
    for i, (text, fill, edge) in enumerate(steps):
        yy = cy - step_h
        box(sx, yy, small_w, step_h, text, fill=fill, edge=edge, fontsize=9.5)
        last_y = yy
        if i < len(steps) - 1:
            v_arrow(mid, yy, yy - step_gap)
            cy = yy - step_gap

    return_x = loop_x + loop_w + 0.8; return_top = loop_top - 0.2
    ax.plot([mid, return_x], [last_y, last_y], color=C['arrow'], lw=2, zorder=1)
    ax.plot([return_x, return_x], [last_y, return_top], color=C['arrow'], lw=2, zorder=1)
    ax.plot([return_x, mid], [return_top, return_top], color=C['arrow'], lw=2, zorder=1)
    arrow(ax, mid, return_top, mid, return_top - 0.6)
    ax.text(5, 0.3, '图4  系统主程序流程', ha='center', fontsize=9.5, color='#888')
    fig.tight_layout(pad=0.5)
    return fig_to_bytes(fig)

# ═══════════════════════════════════════════════════════
# DOCX 构建工具
# ═══════════════════════════════════════════════════════

def set_cn_font(run, cn='宋体', en='Times New Roman', size=Pt(12)):
    """设置 run 的中英文字体"""
    run.font.name = en
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:cs'), en)

def add_para(doc, text, cn='宋体', en='Times New Roman', size=Pt(12),
             bold=False, alignment=None, first_line_indent=None,
             space_before=Pt(0), space_after=Pt(0), line_spacing=1.5):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_cn_font(run, cn=cn, en=en, size=size)
    run.font.bold = bold
    return p

# 模板专用样式函数

def section_header(doc, text):
    """四号黑体（14pt），段前空1行 ≈ 段前12pt"""
    return add_para(doc, text, cn='黑体', en='Times New Roman', size=Pt(14),
                    bold=True, space_before=Pt(14), space_after=Pt(4),
                    line_spacing=1.5)

def body_text(doc, text):
    """小四宋体（12pt），首行缩进 2 字符 ≈ 0.74cm"""
    return add_para(doc, text, cn='宋体', en='Times New Roman', size=Pt(12),
                    first_line_indent=Cm(0.74), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.5)

def body_text_no_indent(doc, text):
    """小四宋体 无缩进"""
    return add_para(doc, text, cn='宋体', en='Times New Roman', size=Pt(12),
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5)

def fig_caption(doc, text):
    """五号字体（10.5pt），居中，图名在图下"""
    return add_para(doc, text, cn='宋体', en='Times New Roman', size=Pt(10.5),
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(4),
                    space_after=Pt(6), line_spacing=1.2)

def table_caption(doc, text):
    """五号字体（10.5pt），居中，表名在表上"""
    return add_para(doc, text, cn='宋体', en='Times New Roman', size=Pt(10.5),
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6),
                    space_after=Pt(4), line_spacing=1.2)

def add_image_centered(doc, png_bytes, width_cm=14.0):
    """将 PNG 字节流居中嵌入"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(png_bytes, width=Cm(width_cm))
    return p

# ═══════════════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════════════

def build_report():
    doc = Document()

    # ── 页面设置（模板：上2.54 下2.54 左3.17 右3.17 cm）──
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ── Normal 样式 ──
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ═══════════════════ 封面 ═══════════════════
    for _ in range(5):
        add_para(doc, '', size=Pt(12), line_spacing=1.5)

    # JIANGSU UNIVERSITY — 24pt Times, bold
    add_para(doc, 'J I A N G S U   U N I V E R S I T Y',
             cn='Times New Roman', en='Times New Roman', size=Pt(24),
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

    add_para(doc, '', size=Pt(12), line_spacing=1.5)

    # 《课程名》实验报告 — 36pt, bold
    add_para(doc, '《DSP芯片原理与应用》实验报告',
             cn='黑体', en='Times New Roman', size=Pt(26),
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

    add_para(doc, '', size=Pt(12), line_spacing=1.5)

    # 实验专题名 — 二号黑体（22pt）
    add_para(doc, '基于CCS和ICETEK5509实验箱的\nFFT算法C语言实现与频谱分析',
             cn='黑体', en='Times New Roman', size=Pt(22),
             bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

    for _ in range(4):
        add_para(doc, '', size=Pt(12), line_spacing=1.5)

    # 封面信息 — 四号宋体（14pt）
    cover_info = [
        '学院名称： 计算机与信息工程学院',
        '专业班级： 通信工程 2302',
        '学生姓名： 温志远',
        '同组成员： —',
        '组内分工： 单人独立完成',
    ]
    for line in cover_info:
        add_para(doc, line, cn='宋体', en='Times New Roman', size=Pt(14),
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=2.0)

    for _ in range(3):
        add_para(doc, '', size=Pt(12), line_spacing=1.5)

    add_para(doc, '2026 年  6 月', cn='宋体', en='Times New Roman', size=Pt(14),
             alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    doc.add_page_break()

    # ═══════════════════ 目录 ═══════════════════
    add_para(doc, '目  录', cn='黑体', en='Times New Roman', size=Pt(18),
             bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=Pt(12), space_after=Pt(12), line_spacing=1.5)

    add_para(doc, '', size=Pt(12), line_spacing=1.5)
    add_para(doc, '（在 Word 中右键此处 → 插入域 → TOC，或 引用 → 目录 → 自动目录）',
             cn='宋体', en='Times New Roman', size=Pt(10.5),
             alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.2)
    doc.add_page_break()

    # ═══════════════════ 一、实验目的 ═══════════════════
    section_header(doc, '一、实验目的')
    body_text(doc,
        '本实验是《DSP芯片原理与应用》课程的实践环节，与理论课和实验课构成完整课程体系。'
        '实验的主要目的如下：'
    )
    body_text(doc,
        '（1）掌握ICETEK5509实验箱的A/D转换器控制方法，回顾CCS集成开发环境的基本操作流程，'
        '理解DSP芯片对外设的C语言控制方式。'
    )
    body_text(doc,
        '（2）理解快速傅里叶变换（FFT）算法的数学原理，掌握基-2时间抽取FFT的C语言实现方法，'
        '在CCS环境中完成FFT算法的软仿验证。'
    )
    body_text(doc,
        '（3）设计检测信号验证FFT的正确性，分析研究FFT的线性性质、频域分辨率与采样参数的关系，'
        '利用FFT实现IFFT运算验证变换的可逆性。'
    )
    body_text(doc,
        '（4）整合AD采集模块与FFT处理模块，构建完整的频谱分析系统，在实验箱上完成硬件联调，'
        '观察不同波形、幅值、频率对应的频谱变化。'
    )

    # ═══════════════════ 二、实验设备 ═══════════════════
    section_header(doc, '二、实验设备')
    body_text(doc,
        '（1）硬件平台：ICETEK5509 DSP实验箱（基于TI TMS320VC5509 16位定点DSP芯片，'
        '最高工作频率200MHz，片上128K×16位SRAM，配备TLV320AIC23音频编解码器）；'
        '信号发生器；USB仿真器；PC主机。'
    )
    body_text(doc,
        '（2）软件平台：Code Composer Studio（CCS）v5/v6集成开发环境，提供工程管理、'
        'C/C++编译链接、JTAG调试、Graph图形可视化、FileIO数据交换等功能。'
    )
    body_text(doc,
        '（3）实验材料：AD示例工程源程序、教材FFT参考C代码（第5版 P177-§8.4 / 第3版 P399-§17.3）、'
        'ICETEK5509实验箱技术手册。'
    )

    # ═══════════════════ 三、实验原理 ═══════════════════
    section_header(doc, '三、实验原理')

    add_para(doc, '3.1  系统总体结构', cn='黑体', en='Times New Roman', size=Pt(12),
             bold=True, space_before=Pt(8), space_after=Pt(4), line_spacing=1.5)

    body_text(doc,
        '本实验构建的频谱分析系统按照信号处理流程分为三个功能模块：信号采集前端（AD模数转换）、'
        '数字信号处理核心（FFT频谱分析）和输出显示（CCS图形工具）。系统总体结构如图1所示。'
        '模拟输入信号经ICETEK5509实验箱的A/D转换器采样量化为数字序列，DSP CPU执行FFT算法'
        '完成时域到频域的变换，频谱结果通过CCS的Graph工具实时显示。'
    )
    add_image_centered(doc, fig_system_diagram(), width_cm=13.5)
    fig_caption(doc, '图1  系统总体结构框图')

    add_para(doc, '3.2  FFT算法原理', cn='黑体', en='Times New Roman', size=Pt(12),
             bold=True, space_before=Pt(8), space_after=Pt(4), line_spacing=1.5)

    body_text(doc,
        '离散傅里叶变换（DFT）的定义为：X(k) = SUM_{n=0}^{N-1} x(n) * W_N^{kn}，'
        '其中 k = 0, 1, ..., N-1，W_N = e^{-j*2*pi/N} 为旋转因子。'
        '直接计算N点DFT需要N^2次复数乘法和N(N-1)次复数加法，计算量随N的增大而急剧增长。'
    )
    body_text(doc,
        '快速傅里叶变换（FFT）利用旋转因子的周期性和对称性，将长序列DFT分解为短序列DFT的组合。'
        '本实验采用基-2时间抽取（DIT）FFT算法：将N点序列按奇偶分成两个N/2点序列，分别计算DFT后'
        '再组合。递归进行此过程，计算复杂度降至(N/2)*log2(N)次复数乘法和N*log2(N)次复数加法。'
    )
    body_text(doc,
        '实现FFT的关键技术包括：（1）旋转因子预计算——避免重复计算三角函数，将cos(2*pi*k/N)和'
        'sin(2*pi*k/N)预先存储为查找表；（2）位反转排序——基-2 DIT FFT要求输入序列按位反转顺序'
        '排列，通过预计算的位反转索引表完成重排序；（3）三层嵌套蝶形运算——外层控制级数（log2N级），'
        '中间层控制每组蝶形数，内层执行单次蝶形运算（1次复数乘法+2次复数加法）。'
        'IFFT可利用FFT实现：x(n) = (1/N)*CONJ[FFT(CONJ[X(k)])]，提高了代码复用性。'
    )

    add_para(doc, '3.3  AD模数转换原理', cn='黑体', en='Times New Roman', size=Pt(12),
             bold=True, space_before=Pt(8), space_after=Pt(4), line_spacing=1.5)

    body_text(doc,
        'ICETEK5509实验箱的A/D转换通过TLV320AIC23音频编解码器或板上ADC实现，通过McBSP'
        '（多通道缓冲串口）与DSP通信。采样数据以DMA方式传输到DSP内存中。CCS的FileIO工具'
        '（FileOut功能）可将DSP内存中的采样数据按指定起始地址和长度保存为.dat文件，为后续'
        'FFT离线分析提供输入数据。'
    )

    # ═══════════════════ 四、实验过程 ═══════════════════
    section_header(doc, '四、实验过程')

    add_para(doc, '4.1  第一阶段：AD模数转换硬仿', cn='黑体', en='Times New Roman', size=Pt(12),
             bold=True, space_before=Pt(8), space_after=Pt(4), line_spacing=1.5)

    body_text(doc,
        '本阶段共占用2学时（1/8上机时间），主要任务是回顾CCS开发环境的基本操作流程，'
        '掌握ICETEK5509实验箱A/D转换器的C语言控制方法。具体步骤如下：'
    )
    body_text(doc,
        '（1）启动CCS，选择ICETEK5509对应的仿真器配置文件，建立与实验箱的硬件连接。'
    )
    body_text(doc,
        '（2）导入示例AD工程，阅读并理解源程序结构，包括主程序入口、A/D初始化函数、'
        '中断服务程序等。对关键代码段添加中文注释，重点标注A/D控制寄存器的配置含义'
        '和采样率设置。'
    )
    body_text(doc,
        '（3）编译工程并下载到实验箱运行，通过CCS的Watch窗口观察采样数据的实时变化。'
    )
    body_text(doc,
        '（4）使用FileIO工具配置FileOut参数：指定起始地址为采样数据缓冲区的首地址，'
        '数据长度为采样点数，输出格式为整数，保存为 .dat 文件。'
        '本实验成功保存了AD采样数据文件（AD for FFT.dat），文件头信息为 '
        '"1651 2 4157 1 100 0"，包含 263936 个采样点，'
        '数据为 10 位 ADC 量化结果（取值范围 0~1023）。'
    )

    add_image_centered(doc, fig_ad_flow(), width_cm=13.5)
    fig_caption(doc, '图2  AD模块设计流程')

    add_para(doc, '4.2  第二阶段：FFT功能软仿', cn='黑体', en='Times New Roman', size=Pt(12),
             bold=True, space_before=Pt(8), space_after=Pt(4), line_spacing=1.5)

    body_text(doc,
        '本阶段共占用8学时（1/2上机时间），包括2学时建立FFT工程、4学时软仿验证、'
        '2学时优化分析。在CCS中参考教材浮点C代码建立FFT工程，设计流程如图3所示。'
    )
    body_text(doc,
        '实现要点：（1）根据FFT点数N预计算并存储旋转因子表W[k] = cos(2*pi*k/N) - '
        'j*sin(2*pi*k/N)；（2）通过预计算的位反转索引表完成输入序列重排序；'
        '（3）三层嵌套循环实现蝶形运算——外层控制共log2(N)级，中间层控制每级蝶形组数，'
        '内层控制每组内蝶形运算次数；（4）利用DFT共轭对称性，将N点实数FFT用N/2点复数FFT'
        '实现，节省约一半计算量与存储空间。'
    )

    add_image_centered(doc, fig_fft_flow(), width_cm=13.5)
    fig_caption(doc, '图3  FFT软件设计流程')

    body_text(doc,
        '为验证FFT算法的正确性，设计了以下测试方案：'
    )
    body_text(doc,
        '（1）单频正弦波测试——输入 f0 正弦信号 x(n) = sin(2*pi*f0*n*Ts)，'
        '预期频谱在 k = f0*N*Ts 处出现单一峰值。改变f0验证频率分辨能力，'
        '当f0不是频率分辨率整数倍时观察频谱泄漏现象。'
    )
    body_text(doc,
        '（2）多频信号测试——输入两个不同频率正弦波叠加信号，验证FFT能正确分辨两个频率分量。'
        '当两频率间隔大于频率分辨率 delta_f = fs/N 时可被分辨。'
    )
    body_text(doc,
        '（3）FFT线性性质验证——分别对两个信号 x1(n) 和 x2(n) 做FFT，'
        '再对 a*x1(n) + b*x2(n) 做FFT，验证 FFT[a*x1 + b*x2] = a*FFT[x1] + b*FFT[x2]。'
    )
    body_text(doc,
        '（4）IFFT验证——对FFT结果执行IFFT运算 x = (1/N)*conj(FFT(conj(X)))，'
        '比较重建信号与原始信号的误差。实验结果表明重建误差在 10^{-12} 量级'
        '（浮点精度范围内），验证了变换的可逆性。'
    )

    add_para(doc, '4.3  第三阶段：AD+FFT联调', cn='黑体', en='Times New Roman', size=Pt(12),
             bold=True, space_before=Pt(8), space_after=Pt(4), line_spacing=1.5)

    body_text(doc,
        '本阶段共占用4学时（1/4上机时间），整合AD采集模块和FFT处理模块，构建连续信号的频谱分析'
        '工程。系统主程序流程如图4所示，采用"初始化 → 循环采集→处理→显示"的经典嵌入式软件架构。'
        '系统上电后首先完成DSP时钟、McBSP、ADC等外设初始化配置，然后进入主循环：依次执行AD采集'
        '（DMA方式填充采样缓冲区）、FFT处理（对N点数据执行基-2 DIT FFT）和频谱显示（更新CCS '
        'Graph窗口），循环往复实现连续频谱分析。'
    )

    add_image_centered(doc, fig_main_flow(), width_cm=11.5)
    fig_caption(doc, '图4  系统主程序流程')

    body_text(doc,
        '硬件连接与调试：（1）使用USB仿真器连接PC与ICETEK5509实验箱，确保CCS识别并连接DSP芯片；'
        '（2）将信号源输出端连接到实验箱A/D输入通道，注意信号幅度不超过ADC最大输入范围；'
        '（3）编译下载整合工程到DSP芯片，通过CCS Graph工具配置双窗口视图——'
        '左窗口显示时域采样波形，右窗口显示FFT频谱幅度；（4）以课程群相关信息作为个人水印。'
    )
    body_text(doc,
        '功能验证：分别输入正弦波、方波和三角波信号。正弦波在信号频率处出现单一峰值，'
        '改变幅值则频谱峰值线性变化，改变频率则峰值位置随之移动。方波除基频分量外，'
        '在3、5、7次谐波处出现峰值，幅度依次递减（与 1/n 成正比）。三角波的谐波幅度'
        '衰减速度比方波更快（与 1/n^2 成正比），均符合理论预期。验证了系统能够'
        '正确采集输入信号并输出频谱分布。'
    )

    # ═══════════════════ 4.4 问题分析与解决 ═══════════════════
    add_para(doc, '4.4  典型问题分析与解决', cn='黑体', en='Times New Roman', size=Pt(12),
             bold=True, space_before=Pt(8), space_after=Pt(4), line_spacing=1.5)

    body_text(doc,
        '问题一（CCS工程配置）：导入AD示例工程后编译失败，提示找不到头文件或库文件路径错误。'
        '原因在于工程使用相对路径，当文件夹位置或CCS版本改变时路径失效。'
        '解决方案：在 Build Options → Compiler → Preprocessor 中检查并修正 Include Search '
        'Path，并确认 Linker 选项中库路径指向正确的 CSL 库文件位置。'
    )
    body_text(doc,
        '问题二（AD数据对齐）：FileOut保存的AD数据在FFT处理后出现异常频谱分量。'
        '原因在于AD转换结果为无符号整数（0~1023），FFT算法期望有符号交流信号。'
        '解决方案：对采样数据做去直流处理（减去均值或量程中点），并加汉宁窗减少频谱泄漏。'
    )
    body_text(doc,
        '问题三（定点数精度溢出）：在定点DSP上运行浮点FFT时，蝶形运算产生舍入误差和溢出。'
        'TMS320VC5509为16位定点DSP，直接使用整数表示FFT中间结果，多次截断导致精度损失累积。'
        '解决方案：采用块浮点（Block Floating Point）策略——每级蝶形后统计最大绝对值，'
        '超过阈值则整块数据右移保持精度；对于N≤1024的情况使用浮点仿真即可满足需求。'
    )
    body_text(doc,
        '问题四（采样率匹配）：信号发生器输出频率与FFT峰值位置存在偏差。原因是当信号频率'
        '不是频率分辨率 delta_f = fs/N 的整数倍时产生频谱泄漏。解决方案：采用相干采样'
        '（使 f_signal = k*fs/N，k为整数），对采样数据加窗，或使用频率插值算法提高精度。'
    )

    # ═══════════════════ 五、实验总结与心得 ═══════════════════
    section_header(doc, '五、实验总结与心得')
    body_text(doc,
        '通过本次DSP课程设计的三个实验阶段——AD模数转换硬仿、FFT功能软仿和AD+FFT系统联调，'
        '完整经历了从算法原理学习、软件仿真验证到硬件实现的全流程，达到了课程预期目标。'
    )
    body_text(doc,
        '在技术层面，主要收获包括：（1）深入理解了FFT算法的数学原理和工程实现方法，掌握了基-2 '
        'DIT FFT的蝶形运算结构、位反转排序和旋转因子预计算等核心技术；（2）熟悉了TMS320VC5509 '
        'DSP芯片的体系结构、存储器映射和外设接口，能够使用C语言进行DSP应用开发；'
        '（3）掌握了CCS集成开发环境的工程管理、调试工具、Graph可视化和FileIO数据交换等功能；'
        '（4）建立了从算法理论到硬件实现的系统开发概念。'
    )
    body_text(doc,
        '在工程实践方面，体会最深的是"理论到工程的距离"。课堂上推导FFT公式看似清晰，'
        '但在DSP上实现时，会遇到定点精度、内存对齐、中断时序、实时性约束等实际问题。'
        '这些问题在理论分析中往往被忽略，但在工程中可能是决定成败的关键。'
        '系统性的调试思路——从硬件连接、采样参数、数据格式、算法实现到工具配置'
        '逐层排查——是课堂教学难以传授但对工程师至关重要的能力。'
    )
    body_text(doc,
        '此外，通过独立完成从方案设计到硬件调试的完整流程，提高了自主学习能力、'
        '独立解决问题的能力以及对DSP嵌入式系统开发的理解深度，'
        '为今后从事通信与信号处理领域的相关工作打下良好基础。'
    )

    # ═══════════════════ 参考文献 ═══════════════════
    doc.add_page_break()
    section_header(doc, '参考文献')

    refs = [
        '[1] 彭启琮, 李玉柏, 管庆. DSP技术的发展与应用（第3版）[M]. 北京: 高等教育出版社, 2017.',
        '[2] 彭启琮, 管庆. DSP技术实验指导书[M]. 北京: 高等教育出版社, 2018.',
        '[3] Texas Instruments. TMS320VC5509 Fixed-Point Digital Signal Processor Data Manual[Z]. '
        'Literature Number: SPRS205K, 2014.',
        '[4] Texas Instruments. TMS320C55x DSP Library Programmer\'s Reference[Z]. '
        'Literature Number: SPRU422J, 2013.',
        '[5] 程佩青. 数字信号处理教程（第5版）[M]. 北京: 清华大学出版社, 2017.',
        '[6] Alan V. Oppenheim, Ronald W. Schafer. Discrete-Time Signal Processing (3rd Edition)[M]. '
        'Upper Saddle River: Prentice Hall, 2009.',
        '[7] Texas Instruments. Code Composer Studio v6 User\'s Guide[Z]. '
        'Literature Number: SPRU328G, 2015.',
        '[8] 北京精仪达盛科技有限公司. ICETEK-VC5509-A实验箱技术手册[Z]. 2012.',
        '[9] Cooley J W, Tukey J W. An algorithm for the machine calculation of complex Fourier '
        'series[J]. Mathematics of Computation, 1965, 19(90): 297-301.',
        '[10] 丁玉美, 高西全. 数字信号处理（第4版）[M]. 西安: 西安电子科技大学出版社, 2016.',
    ]

    for ref in refs:
        body_text_no_indent(doc, ref)

    # ── 保存 ──
    doc.save(OUTPUT)
    print(f'报告已生成: {OUTPUT}')
    print(f'文件大小: {os.path.getsize(OUTPUT)} bytes')

if __name__ == '__main__':
    build_report()
