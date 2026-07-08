#!/usr/bin/env python3
"""Final cleanup: remove duplicate 目录, ensure clean page numbering."""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

INPUT = '/Users/wenzhiyuan/Desktop/ff1/output/DSP课程设计报告_温志远.docx'
OUTPUT = INPUT

doc = Document(INPUT)

# ── Step 1: Find and remove duplicate "目录" ──
# The proper TOC starts at the section break we inserted.
# The stray "目录" is in section 0 (cover area). Let's find and delete it.

to_delete = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # Find the stray 目录 that's in the cover/TOC gap area (before the real TOC)
    # The real TOC has 目  录 (with spaces), the stray one is 目录
    if text == '目录' and i < 25:
        to_delete.append(i)
        print(f"Found stray 目录 at P{i}")

# Remove from bottom to top to preserve indices
for idx in reversed(to_delete):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)
    print(f"  Removed P{idx}")

if not to_delete:
    print("No stray 目录 found — checking what's between cover and TOC...")
    # Let's look at paragraphs 15-28
    for i in range(15, min(40, len(doc.paragraphs))):
        t = doc.paragraphs[i].text.strip()
        if t:
            print(f"  P{i}: '{t[:80]}'")

# ── Step 2: Clean section 0 footer (no number on cover) ──
sec0 = doc.sections[0]
# Already has titlePg=True, that's sufficient
# Make sure footer is truly empty
footer0 = sec0.footer
footer0.is_linked_to_previous = False
for p in footer0.paragraphs:
    for r in p.runs:
        r._element.getparent().remove(r._element)

# ── Step 3: Verify TOC section footer (Roman) ──
sec1 = doc.sections[1]
print(f"\nSection 1 (TOC) footer check:")
footer1 = sec1.footer
footer1.is_linked_to_previous = False
# Already set up from before, just verify
for j, p in enumerate(footer1.paragraphs):
    has_page = 'PAGE' in etree.tostring(p._element, encoding='unicode')
    print(f"  P{j}: PAGE field present = {has_page}")

# ── Step 4: Verify body section footer (Arabic) ──
sec2 = doc.sections[2]
print(f"\nSection 2 (Body) footer check:")
footer2 = sec2.footer
footer2.is_linked_to_previous = False
for j, p in enumerate(footer2.paragraphs):
    has_page = 'PAGE' in etree.tostring(p._element, encoding='unicode')
    print(f"  P{j}: PAGE field present = {has_page}")

# ── Save ──
doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")

# ── Final quick check ──
doc2 = Document(OUTPUT)
total = len(doc2.paragraphs)
print(f"Total paragraphs: {total}")
# Count pages approximately (rough estimate based on content)
# Show first 10 non-empty paragraphs between 12-40
print("\nParagraphs 12-40:")
for i in range(12, min(42, total)):
    t = doc2.paragraphs[i].text.strip()
    prefix = '  ' if t else '(empty)'
    print(f"  P{i}: {prefix} {t[:100] if t else ''}")
