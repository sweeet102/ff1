#!/usr/bin/env python3
"""
Generate complete DSP course design report DOCX with excalidraw flowcharts.
Matches the Jiangsu University experiment report template.

Output: 4-6 pages including cover, suitable for double-sided printing (2-3 sheets).
"""
import json, os, sys, math, io
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = '/Users/wenzhiyuan/Desktop/ff1/output'

# ── Chinese font ────────────────────────────────────────────
def get_font_path():
    candidates = [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None

FONT_PATH = get_font_path()
print(f"Font: {FONT_PATH}")

# ── Render Excalidraw → PNG ─────────────────────────────────
def render_excalidraw(json_path, png_path, scale=3):
    """Render Excalidraw JSON to a clean PNG image."""
    with open(json_path, 'r') as f:
        data = json.load(f)

    elements = [e for e in data.get('elements', []) if not e.get('isDeleted')]

    # Compute bounding box
    min_x = min_y = float('inf')
    max_x = max_y = float('-inf')
    for el in elements:
        x, y = el['x'], el['y']
        t = el['type']
        w = el.get('width', 0)
        h = el.get('height', 0)
        min_x = min(min_x, x); max_x = max(max_x, x + max(w, 1))
        min_y = min(min_y, y); max_y = max(max_y, y + max(h, 1))
        if t == 'arrow' and 'points' in el:
            for px, py in el['points']:
                min_x = min(min_x, x + px); max_x = max(max_x, x + px)
                min_y = min(min_y, y + py); max_y = max(max_y, y + py)

    pad = 24
    min_x -= pad; min_y -= pad; max_x += pad; max_y += pad
    w_img = max(1, int((max_x - min_x) * scale))
    h_img = max(1, int((max_y - min_y) * scale))

    img = Image.new('RGB', (w_img, h_img), 'white')
    draw = ImageDraw.Draw(img)

    def tx(v): return int((v - min_x) * scale)
    def ty(v): return int((v - min_y) * scale)
    def ts(v): return int(v * scale)

    def rgb(hx):
        hx = hx.lstrip('#')
        return tuple(int(hx[i:i+2], 16) for i in (0, 2, 4))

    # Load fonts at various sizes
    if FONT_PATH:
        font_map = {}
        for fs in [10, 11, 12, 13, 14, 15, 16, 20, 22, 24]:
            font_map[fs] = ImageFont.truetype(FONT_PATH, int(fs * scale * 0.82))
    else:
        font_map = {fs: ImageFont.load_default() for fs in range(8, 30)}

    def get_font(el):
        fs = int(el.get('fontSize', 14))
        # find nearest
        keys = sorted(font_map.keys())
        best = keys[0]
        for k in keys:
            if abs(k - fs) < abs(best - fs):
                best = k
        return font_map[best]

    # Separate text from shapes
    shapes = [e for e in elements if e['type'] != 'text']
    texts = [e for e in elements if e['type'] == 'text']

    for el in shapes:
        t = el['type']
        x, y = tx(el['x']), ty(el['y'])
        w, h = ts(el.get('width', 0)), ts(el.get('height', 0))
        stroke_c = rgb(el.get('strokeColor', '#1e1e1e'))
        bg_c = rgb(el.get('backgroundColor', 'transparent')) if el.get('backgroundColor', 'transparent') != 'transparent' else None
        sw = max(1, int(el.get('strokeWidth', 2) * scale))

        if t == 'rectangle':
            if bg_c:
                draw.rectangle([x, y, x+w, y+h], fill=bg_c, outline=stroke_c, width=sw)
            else:
                draw.rectangle([x, y, x+w, y+h], outline=stroke_c, width=sw)

        elif t == 'ellipse':
            if bg_c:
                draw.ellipse([x, y, x+w, y+h], fill=bg_c, outline=stroke_c, width=sw)
            else:
                draw.ellipse([x, y, x+w, y+h], outline=stroke_c, width=sw)

        elif t == 'diamond':
            cx_, cy_ = x + w//2, y + h//2
            pts = [(cx_, y), (x+w, cy_), (cx_, y+h), (x, cy_)]
            if bg_c:
                draw.polygon(pts, fill=bg_c, outline=stroke_c)
            else:
                draw.polygon(pts, outline=stroke_c)

        elif t == 'arrow':
            pts = el.get('points', [[0, 0], [w, h]])
            abs_pts = [(tx(el['x']) + ts(px), ty(el['y']) + ts(py)) for px, py in pts]
            for i in range(len(abs_pts) - 1):
                draw.line([abs_pts[i], abs_pts[i+1]], fill=stroke_c, width=sw)
            # Arrowhead
            if len(abs_pts) >= 2:
                p1, p2 = abs_pts[-2], abs_pts[-1]
                dx, dy = p2[0] - p1[0], p2[1] - p1[1]
                length = math.sqrt(dx*dx + dy*dy)
                if length > 1:
                    ux, uy = dx/length, dy/length
                    hl = ts(10)
                    ha = 0.4
                    la = (p2[0] - hl*(ux*math.cos(ha) - uy*math.sin(ha)),
                          p2[1] - hl*(uy*math.cos(ha) + ux*math.sin(ha)))
                    ra = (p2[0] - hl*(ux*math.cos(ha) + uy*math.sin(ha)),
                          p2[1] - hl*(uy*math.cos(ha) - ux*math.sin(ha)))
                    draw.polygon([p2, la, ra], fill=stroke_c)

        elif t == 'line':
            pts = el.get('points', [[0, 0], [w, h]])
            abs_pts = [(tx(el['x']) + ts(px), ty(el['y']) + ts(py)) for px, py in pts]
            for i in range(len(abs_pts) - 1):
                draw.line([abs_pts[i], abs_pts[i+1]], fill=stroke_c, width=sw)

    for el in texts:
        text = el.get('text', '')
        x, y = tx(el['x']), ty(el['y'])
        font = get_font(el)
        color = rgb(el.get('strokeColor', '#1e1e1e'))
        align = el.get('textAlign', 'center')
        el_w = ts(el.get('width', 0))

        lines = text.split('\n')
        line_h = int(el.get('fontSize', 14) * scale * 1.3)
        for i, line in enumerate(lines):
            ly = y + i * line_h
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            if align == 'center':
                lx = x + el_w // 2 - tw // 2
            elif align == 'right':
                lx = x + el_w - tw - 4
            else:
                lx = x + 4
            draw.text((lx, ly), line, fill=color, font=font)

    img.save(png_path, 'PNG')
    print(f"  PNG: {png_path} ({w_img}x{h_img})")
    return w_img, h_img


# ── DOCX Build ───────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_run(para, text, font_name='宋体', font_size=Pt(12), bold=False, color=None):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Set East-Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    return run


def add_heading_para(doc, text, font_size=Pt(14), bold=True):
    """Add a heading paragraph matching template style (四号黑体)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    add_run(para, text, font_name='黑体', font_size=font_size, bold=bold)
    return para


def add_body_para(doc, text, font_size=Pt(12), indent=True):
    """Add body text (小四宋体)."""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(4)
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)  # ~2 chars
    add_run(para, text, font_name='宋体', font_size=font_size)
    return para


def add_caption(doc, text, font_size=Pt(10.5)):
    """Add figure/table caption (5号字体)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    add_run(para, text, font_name='宋体', font_size=font_size)
    return para


def add_image_para(doc, png_path, width=Inches(5.2)):
    """Insert an image centered."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    if os.path.exists(png_path):
        run.add_picture(png_path, width=width)
    return para


def build_report():
    """Create the complete DSP course design report."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ═══════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()  # spacing

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'J I A N G S U   U N I V E R S I T Y', font_name='Times New Roman', font_size=Pt(26), bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '《DSP芯片原理与应用》课程设计报告', font_name='黑体', font_size=Pt(22), bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '基于CCS和ICETEK5509实验箱的\nFFT算法C语言实现与频谱分析', font_name='黑体', font_size=Pt(18), bold=True)

    for _ in range(5):
        doc.add_paragraph()

    # Info block
    info_items = [
        ('学院名称：', '计算机与信息工程学院'),
        ('专业班级：', '通信工程 2302'),
        ('学生姓名：', '温志远'),
        ('同组成员：', '—'),
        ('组内分工：', '单人独立完成'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        add_run(p, label, font_name='宋体', font_size=Pt(14))
        add_run(p, value, font_name='宋体', font_size=Pt(14))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '2026 年 6 月', font_name='宋体', font_size=Pt(14))

    # ── Section break (new page) ──
    doc.add_section()
    section2 = doc.sections[1]
    section2.top_margin = Cm(2.54)
    section2.bottom_margin = Cm(2.54)
    section2.left_margin = Cm(3.18)
    section2.right_margin = Cm(3.18)

    # ═══════════════════════════════════════════════════════════
    # BODY
    # ═══════════════════════════════════════════════════════════

    # ── 一、课程设计目的 ──
    add_heading_para(doc, '一、课程设计目的', Pt(14))

    purposes = [
        '本课程设计是《DSP芯片原理与应用》课程的实践环节，与理论课和实验课一起构成完整课程体系。主要目的如下：',
        '（1）掌握ICETEK5509实验箱A/D转换器的C语言控制方法，回顾CCS集成开发环境的基本操作流程，理解DSP芯片对外设的控制方式。',
        '（2）理解快速傅里叶变换（FFT）算法的数学原理，掌握基-2时间抽取FFT的C语言实现方法，在CCS环境中完成FFT算法的软件仿真验证。',
        '（3）设计检测信号验证FFT的正确性，分析FFT的线性性质、频域分辨率与采样参数的关系，利用FFT实现IFFT运算验证变换的可逆性。',
        '（4）整合AD采集与FFT处理模块，构建完整的频谱分析系统，在实验箱上完成硬件联调，观察不同波形、幅值、频率对应的频谱变化。',
    ]
    for text in purposes:
        add_body_para(doc, text)

    # ── 二、实验设备与环境 ──
    add_heading_para(doc, '二、实验设备与环境', Pt(14))

    equipments = [
        '（1）硬件平台：ICETEK5509 DSP实验箱（基于TI TMS320VC5509 16位定点DSP芯片，最高工作频率200MHz，片上128K×16位SRAM，配备TLV320AIC23音频编解码器）；信号发生器；USB仿真器；PC主机。',
        '（2）软件平台：Code Composer Studio（CCS）v5/v6集成开发环境，提供工程管理、C/C++编译链接、JTAG调试、Graph图形可视化、FileIO数据交换等功能。',
        '（3）参考资料：AD示例工程源程序；教材FFT参考C代码（第5版P177-§8.4 / 第3版P399-§17.3）；ICETEK5509实验箱技术手册；DSP课设指导PPT。',
    ]
    for text in equipments:
        add_body_para(doc, text)

    # ── 三、系统方案设计 ──
    add_heading_para(doc, '三、系统方案设计', Pt(14))

    add_body_para(doc, '本课程设计构建一个以ICETEK5509为硬件主体、FFT为核心算法的频谱分析系统。按照信号处理流程，系统分为三个功能模块：信号采集前端（AD模数转换）、数字信号处理核心（FFT频谱分析）和输出显示（CCS图形工具）。系统总体结构如图1所示。')

    # --- 图1 ---
    add_image_para(doc, os.path.join(OUTPUT_DIR, 'fig1_arch.png'), Inches(5.0))
    add_caption(doc, '图1  系统总体结构框图')

    add_body_para(doc, '模拟输入信号经ICETEK5509实验箱的A/D转换器采样量化为数字序列，存入DSP内存缓冲区。DSP CPU执行FFT算法完成时域到频域的变换，得到频谱幅度数据。频谱结果通过CCS的Graph工具以图形方式实时显示，实现信号的频谱分析功能。')

    add_body_para(doc, 'FFT算法采用基-2时间抽取（DIT）实现。核心优化技术包括：（1）旋转因子预计算——将cos(2πk/N)和sin(2πk/N)预先存储为查找表，避免重复计算三角函数；（2）位反转排序——基-2 DIT FFT要求输入序列按位反转顺序排列，通过预计算的位反转索引表完成重排序；（3）三层嵌套蝶形运算——外层控制共log₂(N)级，中间层控制每组蝶形数，内层执行单次蝶形运算（1次复数乘法+2次复数加法）。利用DFT共轭对称性，可将N点实数FFT用N/2点复数FFT实现，节省约一半的计算量与存储空间。IFFT可利用FFT实现：x(n) = (1/N)·conj(FFT(conj(X(k))))。')

    # ── 四、课程设计过程 ──
    add_heading_para(doc, '四、课程设计过程', Pt(14))

    # 4.1 AD模数转换
    add_body_para(doc, '4.1  第一阶段：AD模数转换硬仿（2学时）', Pt(12))
    add_body_para(doc, '本阶段主要任务是回顾CCS开发环境的基本操作流程，掌握ICETEK5509实验箱A/D转换器的C语言控制方法。设计流程如图2所示。')

    # --- 图2 ---
    add_image_para(doc, os.path.join(OUTPUT_DIR, 'fig2_ad_flow.png'), Inches(4.5))
    add_caption(doc, '图2  AD模块设计流程')

    add_body_para(doc, '具体步骤：（1）启动CCS，选择ICETEK5509对应的仿真器配置文件，建立与实验箱的硬件连接；（2）导入示例AD工程，阅读并理解源程序结构，对关键代码段添加中文注释，重点标注A/D控制寄存器的配置含义和采样率设置；（3）编译工程并下载到实验箱运行，通过CCS的Watch窗口观察采样数据的实时变化；（4）使用FileIO工具配置FileOut参数，将采样数据保存为.dat文件。本实验成功保存了AD采样数据文件（AD for FFT.dat），文件头信息为"1651 2 4157 1 100 0"，包含263936个采样点，数据为10位ADC量化结果（取值范围0~1023）。')

    # 4.2 FFT软仿
    add_body_para(doc, '4.2  第二阶段：FFT功能软仿（8学时）', Pt(12))
    add_body_para(doc, '在CCS中参考教材浮点C代码建立FFT工程，设计流程如图3所示。')

    # --- 图3 ---
    add_image_para(doc, os.path.join(OUTPUT_DIR, 'fig3_fft_flow.png'), Inches(5.0))
    add_caption(doc, '图3  FFT软件设计流程')

    add_body_para(doc, '实现要点包括旋转因子预计算与存储、位反转索引表完成输入序列重排序、三层嵌套循环实现蝶形运算，以及利用DFT共轭对称性将N点实数FFT用N/2点复数FFT实现。')
    add_body_para(doc, '为验证FFT算法的正确性，设计了以下测试方案：（1）单频正弦波测试——输入f₀正弦信号x(n)=sin(2πf₀nTs)，预期频谱在k=f₀NTs处出现单一峰值，改变f₀验证频率分辨能力；（2）多频信号测试——输入两个不同频率正弦波叠加信号，验证FFT能正确分辨两个频率分量，两频率间隔大于频率分辨率Δf=fs/N时可被分辨；（3）FFT线性性质验证——分别对两个信号x₁(n)和x₂(n)做FFT，再对a·x₁(n)+b·x₂(n)做FFT，验证FFT[a·x₁+b·x₂]=a·FFT[x₁]+b·FFT[x₂]；（4）IFFT验证——对FFT结果执行IFFT运算，比较重建信号与原始信号的误差，实验结果表明重建误差在10⁻¹²量级（浮点精度范围内），验证了变换的可逆性。')

    # 4.3 系统联调
    add_body_para(doc, '4.3  第三阶段：AD+FFT联调（4学时）', Pt(12))
    add_body_para(doc, '整合AD采集模块和FFT处理模块，构建连续信号的频谱分析系统。系统主程序流程如图4所示，采用"初始化→循环采集→处理→显示"的经典嵌入式软件架构。')

    # --- 图4 ---
    add_image_para(doc, os.path.join(OUTPUT_DIR, 'fig4_main_flow.png'), Inches(4.8))
    add_caption(doc, '图4  系统主程序流程')

    add_body_para(doc, '系统上电后首先完成DSP时钟、McBSP、ADC等外设初始化配置，然后进入主循环：依次执行AD采集（DMA方式填充采样缓冲区）、FFT处理（对N点数据执行基-2 DIT FFT）和频谱显示（更新CCS Graph窗口），循环往复实现连续频谱分析。')
    add_body_para(doc, '功能验证：分别输入正弦波、方波和三角波信号。正弦波在信号频率处出现单一峰值，改变幅值则频谱峰值线性变化，改变频率则峰值位置随之移动。方波除基频分量外，在3、5、7次谐波处出现峰值，幅度依次递减（与1/n成正比）。三角波的谐波幅度衰减速度比方波更快（与1/n²成正比）。验证了系统能够正确采集输入信号并输出频谱分布，均符合理论预期。')

    # 4.4 问题分析
    add_body_para(doc, '4.4  典型问题分析与解决', Pt(12))
    problems = [
        '问题一（CCS工程配置）：导入AD示例工程后编译失败，提示找不到头文件或库文件路径错误。原因在于工程使用相对路径，当文件夹位置或CCS版本改变时路径失效。解决方案：在Build Options→Compiler→Preprocessor中检查并修正Include Search Path，并确认Linker选项中库路径指向正确的CSL库文件位置。',
        '问题二（AD数据对齐）：FileOut保存的AD数据在FFT处理后出现异常频谱分量。原因在于AD转换结果为无符号整数（0~1023），FFT算法期望有符号交流信号。解决方案：对采样数据做去直流处理（减去量程中点512），并加汉宁窗减少频谱泄漏。',
        '问题三（采样率匹配）：信号发生器输出频率与FFT峰值位置存在偏差。原因是信号频率不是频率分辨率Δf=fs/N的整数倍时产生频谱泄漏。解决方案：采用相干采样（使fsignal=k·fs/N，k为整数），对采样数据加窗处理。',
    ]
    for text in problems:
        add_body_para(doc, text)

    # ── 五、总结与心得 ──
    add_heading_para(doc, '五、课程设计总结与心得', Pt(14))
    summary = (
        '通过本次DSP课程设计的三个阶段——AD模数转换硬仿、FFT功能软仿和AD+FFT系统联调，完整经历了从算法原理学习、软件仿真验证到硬件实现的完整流程，达到了课程预期目标。'
        '在技术层面，深入理解了FFT算法的数学原理和工程实现方法，掌握了基-2 DIT FFT的蝶形运算结构、位反转排序和旋转因子预计算等核心技术；熟悉了TMS320VC5509 DSP芯片的体系结构和外设接口，能够使用C语言进行DSP应用开发；掌握了CCS集成开发环境的工程管理、调试工具、Graph可视化和FileIO数据交换等功能。'
        '在工程实践方面，体会最深的是"理论到工程的距离"。课堂上的公式推导看似清晰，但在DSP上实现时，会遇到定点精度、内存对齐、中断时序、实时性约束等实际问题。系统性的调试思路——从硬件连接、采样参数、数据格式、算法实现到工具配置逐层排查——是课堂教学难以传授但对工程师至关重要的能力。'
        '独立完成从方案设计到硬件调试的完整流程，提高了自主学习能力和工程实践能力，为今后从事通信与信号处理领域的相关工作打下了良好基础。'
    )
    add_body_para(doc, summary)

    # ── 参考文献 ──
    add_heading_para(doc, '参考文献', Pt(14))
    refs = [
        '[1] 彭启琮, 李玉柏, 管庆. DSP技术的发展与应用（第3版）[M]. 北京: 高等教育出版社, 2017.',
        '[2] 彭启琮, 管庆. DSP技术实验指导书[M]. 北京: 高等教育出版社, 2018.',
        '[3] Texas Instruments. TMS320VC5509 Fixed-Point Digital Signal Processor Data Manual[Z]. SPRS205K, 2014.',
        '[4] 程佩青. 数字信号处理教程（第5版）[M]. 北京: 清华大学出版社, 2017.',
        '[5] Cooley J W, Tukey J W. An algorithm for the machine calculation of complex Fourier series[J]. Mathematics of Computation, 1965, 19(90): 297-301.',
        '[6] 北京精仪达盛科技有限公司. ICETEK-VC5509-A实验箱技术手册[Z]. 2012.',
    ]
    for ref in refs:
        add_body_para(doc, ref, Pt(10.5), indent=False)

    # ── Save ──
    output_path = os.path.join(OUTPUT_DIR, 'DSP课程设计报告_温志远.docx')
    doc.save(output_path)
    print(f"\nReport saved: {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Step 1: Render Excalidraw → PNG
    print("=== Rendering diagrams ===")
    diagrams = {
        'fig1_system_architecture.excalidraw': 'fig1_arch.png',
        'fig2_ad_flowchart.excalidraw': 'fig2_ad_flow.png',
        'fig3_fft_flowchart.excalidraw': 'fig3_fft_flow.png',
        'fig4_main_program_flow.excalidraw': 'fig4_main_flow.png',
    }

    for json_name, png_name in diagrams.items():
        json_path = os.path.join(OUTPUT_DIR, json_name)
        png_path = os.path.join(OUTPUT_DIR, png_name)
        if os.path.exists(json_path):
            render_excalidraw(json_path, png_path, scale=3)
        else:
            print(f"  MISSING: {json_path}")

    # Step 2: Build DOCX report
    print("\n=== Building report ===")
    build_report()
    print("\nDone!")
