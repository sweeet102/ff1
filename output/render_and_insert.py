#!/usr/bin/env python3
"""Render Excalidraw JSON to PNG images using PIL, then insert into DOCX."""
import json, os, sys
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# Try to find a Chinese font
def find_font():
    paths = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
    ]
    for p in paths:
        if os.path.exists(p):
            return p
    return None

FONT_PATH = find_font()
print(f"Using font: {FONT_PATH}")

def render_excalidraw_to_png(json_path, png_path, scale=2):
    """Render an Excalidraw JSON file to a PNG image."""
    with open(json_path, 'r') as f:
        data = json.load(f)

    elements = data.get('elements', [])

    # Calculate bounding box
    min_x = float('inf')
    min_y = float('inf')
    max_x = float('-inf')
    max_y = float('-inf')

    for el in elements:
        if el.get('isDeleted'):
            continue
        t = el['type']
        x, y = el['x'], el['y']
        if t in ('rectangle', 'ellipse', 'diamond', 'arrow', 'line', 'text'):
            w = el.get('width', 0)
            h = el.get('height', 0)
            min_x = min(min_x, x)
            min_y = min(min_y, y)
            max_x = max(max_x, x + max(w, abs(w)))
            max_y = max(max_y, y + max(h, abs(h)))
            # For arrows with multiple points
            if t == 'arrow' and 'points' in el:
                for px, py in el['points']:
                    min_x = min(min_x, x + px)
                    min_y = min(min_y, y + py)
                    max_x = max(max_x, x + px)
                    max_y = max(max_y, y + py)

    # Add padding
    pad = 30
    min_x -= pad
    min_y -= pad
    max_x += pad
    max_y += pad

    w_img = int((max_x - min_x) * scale)
    h_img = int((max_y - min_y) * scale)

    if w_img <= 0 or h_img <= 0:
        print(f"  Invalid dimensions for {json_path}")
        return False

    img = Image.new('RGB', (w_img, h_img), 'white')
    draw = ImageDraw.Draw(img)

    # Load fonts
    try:
        font_lg = ImageFont.truetype(FONT_PATH, 16 * scale)
        font_md = ImageFont.truetype(FONT_PATH, 13 * scale)
        font_sm = ImageFont.truetype(FONT_PATH, 11 * scale)
        font_title = ImageFont.truetype(FONT_PATH, 20 * scale)
    except:
        font_lg = ImageFont.load_default()
        font_md = ImageFont.load_default()
        font_sm = ImageFont.load_default()
        font_title = ImageFont.load_default()

    def tx(val):
        return int((val - min_x) * scale)

    def ty(val):
        return int((val - min_y) * scale)

    def ts(val):
        return int(val * scale)

    def hex_to_rgb(h):
        h = h.lstrip('#')
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    # Sort: render non-text first, then text on top
    non_text = [e for e in elements if e['type'] != 'text' and not e.get('isDeleted')]
    text_els = [e for e in elements if e['type'] == 'text' and not e.get('isDeleted')]

    for el in non_text:
        t = el['type']
        x, y = tx(el['x']), ty(el['y'])
        w = ts(el.get('width', 0))
        h = ts(el.get('height', 0))
        stroke = el.get('strokeColor', '#1e1e1e')
        bg = el.get('backgroundColor', 'transparent')
        fill = el.get('fillStyle', 'solid')
        sw = max(1, int(el.get('strokeWidth', 2) * scale))

        if t == 'rectangle':
            color_bg = hex_to_rgb(bg) if bg and bg != 'transparent' and fill == 'solid' else None
            color_stroke = hex_to_rgb(stroke)
            x2, y2 = x + w, y + h
            if color_bg:
                draw.rectangle([x, y, x2, y2], fill=color_bg, outline=color_stroke, width=sw)
            else:
                draw.rectangle([x, y, x2, y2], outline=color_stroke, width=sw)

        elif t == 'ellipse':
            color_bg = hex_to_rgb(bg) if bg and bg != 'transparent' and fill == 'solid' else None
            color_stroke = hex_to_rgb(stroke)
            x2, y2 = x + w, y + h
            if color_bg:
                draw.ellipse([x, y, x2, y2], fill=color_bg, outline=color_stroke, width=sw)
            else:
                draw.ellipse([x, y, x2, y2], outline=color_stroke, width=sw)

        elif t == 'diamond':
            color_bg = hex_to_rgb(bg) if bg and bg != 'transparent' and fill == 'solid' else None
            color_stroke = hex_to_rgb(stroke)
            cx, cy = x + w//2, y + h//2
            diamond_pts = [(cx, y), (x + w, cy), (cx, y + h), (x, cy)]
            if color_bg:
                draw.polygon(diamond_pts, fill=color_bg, outline=color_stroke)
            else:
                draw.polygon(diamond_pts, outline=color_stroke)

        elif t == 'arrow':
            color_stroke = hex_to_rgb(stroke)
            pts = el.get('points', [[0, 0], [w, h]])
            abs_pts = []
            start_x, start_y = x, y
            for px, py in pts:
                abs_pts.append((start_x + ts(px), start_y + ts(py)))

            # Draw lines
            for i in range(len(abs_pts) - 1):
                draw.line([abs_pts[i], abs_pts[i+1]], fill=color_stroke, width=sw)

            # Draw arrowhead at the last segment
            if len(abs_pts) >= 2:
                p1 = abs_pts[-2]
                p2 = abs_pts[-1]
                import math
                dx = p2[0] - p1[0]
                dy = p2[1] - p1[1]
                length = math.sqrt(dx*dx + dy*dy)
                if length > 0:
                    ux, uy = dx/length, dy/length
                    head_len = ts(12)
                    head_angle = 0.4
                    left_x = p2[0] - head_len * (ux * math.cos(head_angle) - uy * math.sin(head_angle))
                    left_y = p2[1] - head_len * (uy * math.cos(head_angle) + ux * math.sin(head_angle))
                    right_x = p2[0] - head_len * (ux * math.cos(head_angle) + uy * math.sin(head_angle))
                    right_y = p2[1] - head_len * (uy * math.cos(head_angle) - ux * math.sin(head_angle))
                    draw.polygon([p2, (left_x, left_y), (right_x, right_y)], fill=color_stroke)

    # Render text elements
    for el in text_els:
        text = el.get('text', '').replace('\n', '\n')
        x, y = tx(el['x']), ty(el['y'])
        fs = el.get('fontSize', 16)
        font_size = int(fs * scale * 0.85)

        # Choose font
        if fs >= 20:
            font = ImageFont.truetype(FONT_PATH, font_size) if FONT_PATH else font_title
        elif fs >= 16:
            font = ImageFont.truetype(FONT_PATH, font_size) if FONT_PATH else font_lg
        elif fs >= 13:
            font = ImageFont.truetype(FONT_PATH, font_size) if FONT_PATH else font_md
        else:
            font = ImageFont.truetype(FONT_PATH, font_size) if FONT_PATH else font_sm

        color = hex_to_rgb(el.get('strokeColor', '#1e1e1e'))
        align = el.get('textAlign', 'center')

        lines = text.split('\n')
        line_height = int(font_size * 1.3)

        for i, line in enumerate(lines):
            ly = y + i * line_height
            if align == 'center':
                # Approximate text width for centering
                try:
                    bbox = draw.textbbox((0, 0), line, font=font)
                    tw = bbox[2] - bbox[0]
                except:
                    tw = len(line) * font_size * 0.5
                lx = x + ts(el.get('width', 0)) // 2 - tw // 2
            elif align == 'right':
                try:
                    bbox = draw.textbbox((0, 0), line, font=font)
                    tw = bbox[2] - bbox[0]
                except:
                    tw = len(line) * font_size * 0.5
                lx = x + ts(el.get('width', 0)) - tw - 5
            else:
                lx = x + 5
            draw.text((lx, ly), line, fill=color, font=font)

    img.save(png_path, 'PNG')
    print(f"  Rendered: {png_path} ({w_img}x{h_img})")
    return True


def insert_images_to_docx(docx_template, docx_output, image_map):
    """
    Insert PNG images back into the document at the original image positions.
    image_map: list of dicts with {paragraph_index, png_path, caption_text}
    """
    doc = Document(docx_template)

    # We need to insert images after the paragraphs that originally had them.
    # The paragraphs are: 40, 55, 60, 69
    img_paragraphs = [40, 55, 60, 69]

    for i, (para_idx, png_path) in enumerate(zip(img_paragraphs, sorted(image_map))):
        # Find the caption paragraph (the one after the image)
        # Actually the image was in paragraph 40, 55, 60, 69
        # Those paragraphs now have the text but no image
        # We should insert the new image before the caption paragraph

        # The original structure: text_para (with image inline) → caption_para
        # Now: text_para (text only) → caption_para
        # We want: text_para → new_image → caption_para → ...

        # Insert new image paragraph after para_idx
        if para_idx < len(doc.paragraphs):
            ref_para = doc.paragraphs[para_idx]
            # Create a new paragraph for the image
            new_para = doc.add_paragraph()
            # Move it after ref_para
            ref_para._element.addnext(new_para._element)

            run = new_para.add_run()
            # Add picture as inline shape
            img_path = os.path.join(OUTPUT_DIR, png_path)
            if os.path.exists(img_path):
                run.add_picture(img_path, width=Inches(5.5))
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_para.paragraph_format.space_before = Pt(6)
                new_para.paragraph_format.space_after = Pt(6)
                print(f"  Inserted {png_path} after paragraph {para_idx}")

    doc.save(docx_output)
    print(f"Saved: {docx_output}")


if __name__ == '__main__':
    # Step 1: Render all excalidraw files to PNG
    diagrams = [
        'fig1_system_architecture.excalidraw',
        'fig2_ad_flowchart.excalidraw',
        'fig3_fft_flowchart.excalidraw',
        'fig4_main_program_flow.excalidraw',
    ]

    png_files = []
    for d in diagrams:
        json_path = os.path.join(OUTPUT_DIR, d)
        png_name = d.replace('.excalidraw', '.png')
        png_path = os.path.join(OUTPUT_DIR, png_name)
        if render_excalidraw_to_png(json_path, png_path):
            png_files.append(png_name)

    print(f"\nRendered {len(png_files)} PNG files")

    # Step 2: Insert into DOCX
    docx_template = os.path.join(OUTPUT_DIR, 'DSP实验报告_温志远_no_images.docx')
    docx_output = os.path.join(OUTPUT_DIR, 'DSP实验报告_温志远_final.docx')

    if os.path.exists(docx_template):
        image_map = sorted(png_files)
        insert_images_to_docx(docx_template, docx_output, image_map)
    else:
        print(f"Template not found: {docx_template}")
