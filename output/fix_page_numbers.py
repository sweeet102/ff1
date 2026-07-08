#!/usr/bin/env python3
"""
Rebuild DOCX with proper 3-section page numbering:
  Section 0: Cover page — no page number
  Section 1: TOC (目录) — Roman numerals (i, ii, ...)
  Section 2: Body — simple Arabic (1, 2, 3, ...)
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
from lxml import etree
import os

INPUT = '/Users/wenzhiyuan/Desktop/ff1/output/DSP课程设计报告_温志远.docx'
OUTPUT = '/Users/wenzhiyuan/Desktop/ff1/output/DSP课程设计报告_温志远.docx'

doc = Document(INPUT)

# ──────────────────────────────────────────────
# Helper functions
# ──────────────────────────────────────────────

def add_run(doc_or_para, text, font_name='宋体', font_size=Pt(12), bold=False):
    """Add a run with EA font set."""
    if hasattr(doc_or_para, 'add_run'):
        para = doc_or_para
        run = para.add_run(text)
    else:
        run = doc_or_para  # already a run
    run.font.size = font_size
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    if font_name != 'Times New Roman':
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    return run

def create_page_number_footer(section, fmt='decimal', start=1):
    """Set up footer with just the page number."""
    sectPr = section._sectPr

    # Set page number type
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))

    # Build footer with PAGE field
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    # Add PAGE field
    run_page = para.add_run()
    run_page.font.size = Pt(10.5)
    rPr = run_page._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.insert(0, rFonts)

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run_page._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run_page._r.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fld_end)

def set_title_page(section):
    """Enable 'different first page' so first page shows no footer."""
    sectPr = section._sectPr
    titlePg = OxmlElement('w:titlePg')
    sectPr.insert(0, titlePg)

# ──────────────────────────────────────────────
# Step 1: Understand current structure
# ──────────────────────────────────────────────
print(f"Document has {len(doc.sections)} sections")
print(f"Document has {len(doc.paragraphs)} paragraphs")

# Find the paragraph that starts the body (after cover info)
body_start_idx = None
for i, p in enumerate(doc.paragraphs):
    if '一、课程设计目的' in p.text:
        body_start_idx = i
        print(f"Body starts at paragraph {i}: '{p.text[:50]}'")
        break

# The plan:
# - Section 0 (cover) currently holds paragraphs 0..body_start_idx-1
# - Section 1 (body) holds paragraphs body_start_idx..end
# We need section 0 = cover, section 1 = TOC, section 2 = body
# But simpler: keep section 0 as cover + blank TOC page, section 1 as body

# Actually the simplest approach:
# Section 0: Cover (no page number)
# Add a TOC page by inserting a section break after cover
# Section 1: TOC (Roman numerals)
# Section 2: Body (Arabic numerals)
# But the body already starts at section 1...

# Let me restructure:
# Section 0 stays as cover (title page, no number)
# Insert a new section between cover and body for TOC
# Section 1 becomes TOC (Roman)
# Section 2 becomes body (Arabic)

# First, let's identify the cover paragraphs
# Cover ends at the paragraph with "2026 年 6 月"
cover_end_idx = None
for i, p in enumerate(doc.paragraphs):
    if '2026 年 6 月' in p.text and i < (body_start_idx or 999):
        cover_end_idx = i
        print(f"Cover ends at paragraph {i}: '{p.text[:50]}'")

# ──────────────────────────────────────────────
# Step 2: Insert section break after cover + TOC
# ──────────────────────────────────────────────
# We'll insert a TOC section by:
# 1. Adding a section break after the cover content
# 2. That creates a new section for TOC

# Add TOC page after cover (before the body)
# Insert a section break (new page) after the cover's last paragraph
cover_last_para = doc.paragraphs[cover_end_idx]

# Create a new paragraph for the section break
new_para = doc.add_paragraph()
# Actually we need to insert the section break at a specific position in the element tree

# Alternative approach: use XML manipulation
body = doc.element.body

# Find all <w:p> elements and the section break between cover and body
all_paras = body.findall(qn('w:p'))
all_sectPrs = body.findall(qn('w:sectPr'))

print(f"XML: {len(all_paras)} paragraphs, {len(all_sectPrs)} section properties")

# The document order in body is:
#   <w:p> ... cover paragraphs ... </w:p>
#   <w:p> ... body paragraphs ... </w:p>
#   <w:sectPr> section 0 properties </w:sectPr>
#   <w:sectPr> section 1 properties </w:sectPr>

# Wait, section breaks in python-docx are stored differently.
# Let me check the actual XML structure.

# In python-docx, each section has its own sectPr.
# Section 0's sectPr is at the end of all its paragraphs (before section 1 starts)
# Section 1's sectPr is at the very end of the body

# To add a TOC section between cover and body:
# 1. Insert a <w:p> with <w:sectPr> inside it after the cover paragraphs
# This creates a new section

# Let's do this at the XML level

para_elements = body.findall(qn('w:p'))
sectPr_elements = body.findall(qn('w:sectPr'))

print(f"Found {len(para_elements)} w:p elements")
print(f"Found {len(sectPr_elements)} w:sectPr elements")

# The last sectPr is the final section's properties (section 1)
# The second-to-last sectPr is section 0's properties (if present)

# Let me find where the cover ends in the XML element tree
cover_last_xml_idx = cover_end_idx  # This maps to the paragraph index

# ──────────────────────────────────────────────
# Step 3: Insert TOC section via XML
# ──────────────────────────────────────────────

# We need to:
# 1. Create TOC content paragraph(s)
# 2. Insert a section break (w:p with w:sectPr child) after cover

# Build the TOC page paragraph
toc_title_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:pPr>
    <w:jc w:val="center"/>
    <w:spacing w:before="240" w:after="120"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:ascii="黑体" w:hAnsi="黑体" w:eastAsia="黑体"/>
      <w:b/>
      <w:sz w:val="32"/>
    </w:rPr>
    <w:t>目  录</w:t>
  </w:r>
</w:p>'''

toc_hint_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:pPr>
    <w:spacing w:line="360" w:lineRule="auto"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>
      <w:sz w:val="24"/>
      <w:color w:val="808080"/>
    </w:rPr>
    <w:t>（在 Word 中：引用 → 目录 → 自动目录，即可生成）</w:t>
  </w:r>
</w:p>'''

toc_content_items = [
    ('一、课程设计目的', ''),
    ('二、实验设备与环境', ''),
    ('三、系统方案设计', ''),
    ('四、课程设计过程', ''),
    ('    4.1  第一阶段：AD模数转换硬仿', ''),
    ('    4.2  第二阶段：FFT功能软仿', ''),
    ('    4.3  第三阶段：AD+FFT联调', ''),
    ('    4.4  典型问题分析与解决', ''),
    ('五、课程设计总结与心得', ''),
    ('参考文献', ''),
]

toc_items_xml = ''
for title, _ in toc_content_items:
    toc_items_xml += f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:pPr>
    <w:spacing w:line="360" w:lineRule="auto"/>
    <w:ind w:left="420" w:hanging="420"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>
      <w:sz w:val="24"/>
    </w:rPr>
    <w:t>{title}</w:t>
  </w:r>
</w:p>'''

# Build the section break paragraph (page break + new section properties)
# This goes at the end of the TOC content
section_break_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:pPr>
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1440" w:right="1800" w:bottom="1440" w:left="1800"/>
      <w:pgNumType w:fmt="upperRoman" w:start="1"/>
    </w:sectPr>
  </w:pPr>
</w:p>'''

# Combine all TOC content
toc_all_xml = toc_title_xml + '\n' + toc_hint_xml + '\n' + toc_items_xml + '\n' + section_break_xml

# Now insert the TOC before the body starts
# The body starts at paragraph body_start_idx in python-docx
# In the XML, this corresponds to para_elements[body_start_idx]

# Insert TOC content right before the body's first paragraph
body_first_xml = para_elements[body_start_idx]
toc_fragment = etree.fromstring(f'<root xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{toc_all_xml}</root>')

# Insert TOC elements before the body starts
parent = body_first_xml.getparent()
insert_position = parent.index(body_first_xml)

for child in toc_fragment:
    parent.insert(insert_position, child)
    insert_position += 1

# ──────────────────────────────────────────────
# Step 4: Fix section properties
# ──────────────────────────────────────────────

# Now we should have 3 sections:
# Section 0: Cover (original section 0)
# Section 1: TOC (we just inserted via section break)
# Section 2: Body (original section 1)

# Remove old section properties and fix them
# Let's get all sectPr elements again
all_sectPrs_after = body.findall(qn('w:sectPr'))
print(f"After insert: {len(all_sectPrs_after)} sectPr elements")

# We need to fix:
# - Section 0 (cover): titlePg (no number on first page), Roman format not visible
# - Section 1 (TOC): Roman numerals (already set in the inserted section break)
# - Section 2 (body): Arabic numerals, simple format

# The last sectPr should be the body section
# Let's find and fix the body section's properties

# Actually, let's just reload the document to work with python-docx's section model
import tempfile
tmp_path = os.path.join(os.path.dirname(INPUT), '_tmp_rebuild.docx')
doc.save(tmp_path)
doc2 = Document(tmp_path)

print(f"\nAfter rebuild: {len(doc2.sections)} sections")

for i, sec in enumerate(doc2.sections):
    sectPr = sec._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    titlePg = sectPr.find(qn('w:titlePg'))
    fmt = pgNumType.get(qn('w:fmt')) if pgNumType is not None else 'NOT SET'
    start = pgNumType.get(qn('w:start')) if pgNumType is not None else 'N/A'
    print(f"  Section {i}: fmt={fmt}, start={start}, titlePg={titlePg is not None}")

# Fix section 0 (cover): title page, no visible number
sec0 = doc2.sections[0]
set_title_page(sec0)

# Fix section 1 (TOC): Roman numerals
sec1 = doc2.sections[1]
# Make sure fmt is upperRoman
sectPr1 = sec1._sectPr
pgNumType1 = sectPr1.find(qn('w:pgNumType'))
if pgNumType1 is None:
    pgNumType1 = OxmlElement('w:pgNumType')
    sectPr1.append(pgNumType1)
pgNumType1.set(qn('w:fmt'), 'upperRoman')
pgNumType1.set(qn('w:start'), '1')

# Add page number to TOC footer
footer1 = sec1.footer
footer1.is_linked_to_previous = False
for p in footer1.paragraphs:
    p.clear()
para1 = footer1.paragraphs[0] if footer1.paragraphs else footer1.add_paragraph()
para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
para1.paragraph_format.space_before = Pt(0)
para1.paragraph_format.space_after = Pt(0)

# Add PAGE field with Roman numeral
run1 = para1.add_run()
run1.font.size = Pt(10.5)
rPr = run1._r.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rPr.insert(0, rFonts)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
run1._r.append(fld_begin)
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = ' PAGE '
run1._r.append(instr)
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
run1._r.append(fld_end)

# Fix section 2 (body): simple Arabic numbers
sec2 = doc2.sections[2]
sectPr2 = sec2._sectPr
pgNumType2 = sectPr2.find(qn('w:pgNumType'))
if pgNumType2 is None:
    pgNumType2 = OxmlElement('w:pgNumType')
    sectPr2.append(pgNumType2)
pgNumType2.set(qn('w:fmt'), 'decimal')
pgNumType2.set(qn('w:start'), '1')

# Replace body footer with simple number
footer2 = sec2.footer
footer2.is_linked_to_previous = False
for p in footer2.paragraphs:
    p.clear()
para2 = footer2.paragraphs[0] if footer2.paragraphs else footer2.add_paragraph()
para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
para2.paragraph_format.space_before = Pt(0)
para2.paragraph_format.space_after = Pt(0)

run2 = para2.add_run()
run2.font.size = Pt(10.5)
rPr2 = run2._r.get_or_add_rPr()
rFonts2 = OxmlElement('w:rFonts')
rFonts2.set(qn('w:ascii'), 'Times New Roman')
rFonts2.set(qn('w:hAnsi'), 'Times New Roman')
rPr2.insert(0, rFonts2)
fld_begin2 = OxmlElement('w:fldChar')
fld_begin2.set(qn('w:fldCharType'), 'begin')
run2._r.append(fld_begin2)
instr2 = OxmlElement('w:instrText')
instr2.set(qn('xml:space'), 'preserve')
instr2.text = ' PAGE '
run2._r.append(instr2)
fld_end2 = OxmlElement('w:fldChar')
fld_end2.set(qn('w:fldCharType'), 'end')
run2._r.append(fld_end2)

# ──────────────────────────────────────────────
# Step 5: Save
# ──────────────────────────────────────────────
doc2.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")

# Verify
doc3 = Document(OUTPUT)
print(f"\nFinal verification: {len(doc3.sections)} sections")
for i, sec in enumerate(doc3.sections):
    sectPr = sec._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    titlePg = sectPr.find(qn('w:titlePg'))
    fmt = pgNumType.get(qn('w:fmt')) if pgNumType is not None else 'NOT SET'
    start = pgNumType.get(qn('w:start')) if pgNumType is not None else 'N/A'
    footer_text = ''
    if sec.footer and sec.footer.paragraphs:
        footer_text = sec.footer.paragraphs[0].text[:50]
    print(f"  Sec {i}: fmt={fmt}, start={start}, titlePg={titlePg is not None}, footer='{footer_text}'")

print("\nDone! Summary:")
print("  Cover    → no page number")
print("  TOC      → I, II, III, ...")
print("  Body     → 1, 2, 3, ...")

# Clean up temp file
if os.path.exists(tmp_path):
    os.remove(tmp_path)
