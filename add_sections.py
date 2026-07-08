#!/usr/bin/env python3
"""为论文添加分节符、页码和页眉 — 简化可靠版"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DST = '/Users/wenzhiyuan/Desktop/3231303014-王奇坤_formatted.docx'
doc = Document(DST)

# ============================================================
# 1. Find first chapter paragraph and insert section break
# ============================================================
body = doc.element.body
all_paras = list(body.iterchildren(qn('w:p')))

# Find Heading 1 paragraph
chap_idx = None
for i, p in enumerate(all_paras):
    pPr = p.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
            chap_idx = i
            break

if chap_idx is None:
    for i, p in enumerate(all_paras):
        text = ''.join(p.itertext())
        if ('第一章' in text or '第1章' in text) and i > 10:
            chap_idx = i
            break

print(f"第一章位于元素索引: {chap_idx}")

if chap_idx and chap_idx > 0:
    # Insert section break: add w:sectPr to the PREVIOUS paragraph's w:pPr
    prev_p = all_paras[chap_idx - 1]

    pPr = prev_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        prev_p.insert(0, pPr)

    # Create section break properties
    new_sectPr = OxmlElement('w:sectPr')

    # Copy page dimensions etc from current section
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')   # A4 width in twips (21cm)
    pgSz.set(qn('w:h'), '16838')   # A4 height in twips (29.7cm)
    new_sectPr.append(pgSz)

    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1440')     # 2.54cm
    pgMar.set(qn('w:bottom'), '1440')  # 2.54cm
    pgMar.set(qn('w:left'), '1260')    # 2.22cm
    pgMar.set(qn('w:right'), '1106')   # 1.95cm
    pgMar.set(qn('w:header'), '851')   # 1.5cm
    pgMar.set(qn('w:footer'), '992')   # 1.75cm
    new_sectPr.append(pgMar)

    # Page number type: Arabic 1,2,3... for body section
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    new_sectPr.append(pgNumType)

    pPr.append(new_sectPr)
    print("分节符已插入到第一章之前")

# ============================================================
# 2. Update original section (front matter) page number type
# ============================================================
original_sectPr = body.find(qn('w:sectPr'))
if original_sectPr is not None:
    # Set Roman numerals for front matter
    pgNumType = original_sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        original_sectPr.insert(0, pgNumType)
    pgNumType.set(qn('w:fmt'), 'upperRoman')
    pgNumType.set(qn('w:start'), '1')
    print("前置部分页码设为罗马数字")

# ============================================================
# 3. Set up headers and footers via python-docx sections
# ============================================================
print(f"文档现在有 {len(doc.sections)} 个 section")

# Section 1: Front matter (TOC)
sec1 = doc.sections[0]
sec1.different_first_page_header_footer = True  # Cover page no header/footer

# --- Section 1 Footer (Roman numerals) ---
footer1 = sec1.footer
footer1.is_linked_to_previous = False
for p in footer1.paragraphs:
    p.clear()
p = footer1.paragraphs[0] if footer1.paragraphs else footer1.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
# Insert PAGE field using fldChar
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
run._element.append(fld_begin)
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = ' PAGE '
run._element.append(instr)
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
run._element.append(fld_end)

# Header for section 1: 左侧"目录" + 右侧"江苏大学本科毕业设计（论文）"
header1 = sec1.header
header1.is_linked_to_previous = False
for p in header1.paragraphs:
    p.clear()
h1p = header1.paragraphs[0] if header1.paragraphs else header1.add_paragraph()
# Just set right-aligned tab for now
h1p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = h1p.add_run('目录')
r1.font.size = Pt(9)
r1.font.name = '宋体'
# Add right-aligned tab stop
h1p.paragraph_format.tab_stops.add_tab_stop(Cm(14.6), 2)  # RIGHT align
r2 = h1p.add_run('\t江苏大学本科毕业设计（论文）')
r2.font.size = Pt(9)
r2.font.name = '宋体'

print("Section 1 页脚/页眉已设置")

# Section 2: Body (if exists)
if len(doc.sections) >= 2:
    sec2 = doc.sections[1]
    sec2.different_first_page_header_footer = False

    # Footer: Arabic numerals
    footer2 = sec2.footer
    footer2.is_linked_to_previous = False
    for p in footer2.paragraphs:
        p.clear()
    p2 = footer2.paragraphs[0] if footer2.paragraphs else footer2.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run()
    run2.font.size = Pt(9)
    run2.font.name = 'Times New Roman'
    fld_begin2 = OxmlElement('w:fldChar')
    fld_begin2.set(qn('w:fldCharType'), 'begin')
    run2._element.append(fld_begin2)
    instr2 = OxmlElement('w:instrText')
    instr2.set(qn('xml:space'), 'preserve')
    instr2.text = ' PAGE '
    run2._element.append(instr2)
    fld_end2 = OxmlElement('w:fldChar')
    fld_end2.set(qn('w:fldCharType'), 'end')
    run2._element.append(fld_end2)

    # Header: only right side "江苏大学本科毕业设计（论文）"
    header2 = sec2.header
    header2.is_linked_to_previous = False
    for p in header2.paragraphs:
        p.clear()
    h2p = header2.paragraphs[0] if header2.paragraphs else header2.add_paragraph()
    h2p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rh = h2p.add_run('江苏大学本科毕业设计（论文）')
    rh.font.size = Pt(9)
    rh.font.name = '宋体'

    print("Section 2 页脚/页眉已设置")

# ============================================================
# Save
# ============================================================
doc.save(DST)
print(f"\n✅ 分节、页码、页眉设置完成: {DST}")
print("\n⚠️  提醒:")
print("  1. 在 Word 中右键目录 → 更新域（更新整个目录）")
print("  2. 检查页码是否正确显示")
print("  3. 每章应另起一页（在 Word 中每章标题前 Ctrl+Enter 插入分页符）")
print("  4. 指导教师姓名/职称请在封面表格中手动填写")
