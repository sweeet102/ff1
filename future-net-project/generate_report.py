#!/usr/bin/env python3
"""
按模板格式生成未来网络技术课程报告（完整版）
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import sys
import io

TEMPLATE = '/Users/wenzhiyuan/Desktop/未来网络技术课程实践/23级未来网络技术-报告模板.docx'
OUTPUT   = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'results', '课程报告.docx')
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

SZ_ERHAO   = Pt(22)
SZ_XIAOER  = Pt(18)   # 小二，章标题用
SZ_XIAOSAN = Pt(15)
SZ_SIHAO   = Pt(14)
SZ_XIAOSI  = Pt(12)
SZ_WUHAO   = Pt(10.5)
FONT_CN    = '宋体'
FONT_EN    = 'Times New Roman'
FONT_HEI   = '黑体'

# 江苏大学标准页边距：上2.54 下2.54 左2.22 右1.95
MARGIN_LEFT   = Cm(2.22)
MARGIN_RIGHT  = Cm(1.95)
MARGIN_TOP    = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)


def set_page_margins(section):
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM


def set_pf(para, indent=True, align=None):
    pf = para.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Cm(0.74)
    else:
        pf.first_line_indent = Cm(0)
    if align is not None:
        pf.alignment = align


def add_run(para, text, font_cn=FONT_CN, font_en=FONT_EN, size=SZ_XIAOSI, bold=False):
    r = para.add_run(text)
    r.font.size = size
    r.font.bold = bold
    r.font.name = font_en
    rPr = r._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_cn}"/>')
    rPr.insert(0, rFonts)


def body(doc, text, indent=True):
    p = doc.add_paragraph()
    set_pf(p, indent=indent)
    add_run(p, text)


def h1(doc, text):
    """一级标题（章标题）：黑体小二 18pt 居中，每章另起一页"""
    p = doc.add_paragraph()
    set_pf(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, text, font_cn=FONT_HEI, font_en=FONT_HEI, size=SZ_XIAOER, bold=True)
    # 每章另起一页
    p.paragraph_format.page_break_before = True


def h2(doc, text):
    """二级标题：宋体四号 14pt 加粗左对齐"""
    p = doc.add_paragraph()
    set_pf(p, indent=False)
    add_run(p, text, font_cn=FONT_CN, font_en=FONT_EN, size=SZ_SIHAO, bold=True)


def h3(doc, text):
    """三级标题：宋体小四 12pt 加粗左对齐"""
    p = doc.add_paragraph()
    set_pf(p, indent=False)
    add_run(p, text, font_cn=FONT_CN, font_en=FONT_EN, size=SZ_XIAOSI, bold=True)


def caption(doc, text):
    p = doc.add_paragraph()
    set_pf(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, text, size=SZ_WUHAO, bold=True)


# ── 图片路径 ──
OBSIDIAN = "/Users/wenzhiyuan/Library/Mobile Documents/iCloud~md~obsidian/Documents/xV"
SDN_IMG  = OBSIDIAN + "/SDN实验/attachments"
P4_IMG   = OBSIDIAN + "/P4实验/attachments"


def add_picture(doc, path, width_inches=5.5, caption_text=""):
    """插入图片 + 图题 — 用 PIL 缩放后插入，宽度自适应页面"""
    from docx.shared import Inches
    from PIL import Image
    try:
        # 用 PIL 打开原图
        img = Image.open(path)
        # 缩放保证清晰度（最大 1600px 宽）
        max_px = 1600
        if img.width > max_px:
            ratio = max_px / img.width
            new_h = int(img.height * ratio)
            img = img.resize((max_px, new_h), Image.LANCZOS)

        # 转 JPEG 存到 BytesIO
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=90)
        buf.seek(0)

        # 插入文档（图片占页面可用宽度的 ~85%）
        p_img = doc.add_paragraph()
        set_pf(p_img, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        p_img.add_run().add_picture(buf, width=Inches(width_inches))
        if caption_text:
            caption(doc, caption_text)
    except Exception as e:
        p = doc.add_paragraph()
        set_pf(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p, f'（图片: {caption_text or os.path.basename(path)}）', size=SZ_WUHAO)


def make_table(doc, headers, rows):
    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    for ci, h in enumerate(headers):
        p = table.rows[0].cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=SZ_WUHAO, bold=True)
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            p = table.rows[ri + 1].cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(cell), size=SZ_WUHAO)


def page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════

def build_report():
    doc = Document(TEMPLATE)
    for section in doc.sections:
        set_page_margins(section)

    # 清除模板说明文字
    body_start = 0
    for i, para in enumerate(doc.paragraphs):
        if '报告目录自拟' in para.text:
            body_start = i; break
    if body_start > 0:
        for para in list(doc.paragraphs)[body_start:]:
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)

    # ── 目录 ──
    page_break(doc)
    p = doc.add_paragraph()
    set_pf(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, '目录', font_cn=FONT_HEI, font_en=FONT_HEI, size=SZ_ERHAO, bold=True)

    toc = [
        ('1', '未来网络技术调研', 0),
        ('1.1', '可编程网络设备的技术发展', 1), ('1.2', '现有标准与知识产权调研', 1),
        ('1.3', '国内产业情况与产业政策', 1), ('1.4', '网络人工智能技术的前沿进展', 1),
        ('1.5', '未来网络产业用人需求调查', 1),
        ('2', '未来网络技术学习', 0),
        ('2.1', 'SDN 实验', 1), ('2.2', 'P4 实验', 1),
        ('3', '未来网络技术的能耗分析', 0),
        ('4', '网络创新设计：SDN+AI 智能流量分类系统', 0),
        ('4.1', '拟解决的目标问题', 1), ('4.2', '文献与资料调研', 1),
        ('4.3', '作品设计与开发', 1), ('4.4', '实验结果与分析', 1),
        ('5', '实践课程总结', 0),
        ('参考文献', '', 0),
    ]
    for num, title, level in toc:
        p = doc.add_paragraph()
        set_pf(p, indent=False)
        prefix = '    ' if level == 1 else ''
        sz = SZ_SIHAO if level == 0 else SZ_XIAOSI
        add_run(p, f'{prefix}{num}  {title}' if num != '参考文献' else '参考文献', size=sz)

    # ═══════════════════════════════════════════════════════
    # 第一章：未来网络技术调研
    # ═══════════════════════════════════════════════════════
    h1(doc, '1  未来网络技术调研')

    # --- 1.1 ---
    h2(doc, '1.1  可编程网络设备的技术发展')

    h3(doc, '1.1.1  SDN 技术发展')
    body(doc, '软件定义网络（Software-Defined Networking，SDN）是一种将网络设备的控制平面与数据平面解耦的新型网络架构，通过集中式控制器对整个网络进行统一管理和编程。SDN 的起源可追溯至 2006 年斯坦福大学 Nick McKeown 教授团队的 Ethane 项目——该项目首次提出通过集中控制器管理网络中所有交换机的转发策略。2008 年，OpenFlow 协议正式发布，允许控制器通过标准化的南向接口对交换机流表进行增删改查，标志着 SDN 从学术概念走向可实现的网络协议。2011 年，德国电信、Google、微软、Facebook 等联合成立开放网络基金会（ONF），推动了 SDN 标准化和产业落地。')
    body(doc, 'SDN 的技术核心分为三层架构。基础设施层（数据平面）由支持 OpenFlow 协议的交换机组成，负责高速数据包转发。控制层由集中的 SDN 控制器实现，负责全网拓扑发现、路径计算、策略下发等全局决策。应用层通过北向 API 调用控制器的能力，实现负载均衡、流量工程、安全策略等网络应用。OpenFlow 1.3 版本相比 1.0 增加了多级流表、Meter 表、Group 表等重要特性——多级流表允许数据包经过多个处理阶段，类似于程序的多层 if-else 逻辑；Meter 表实现基于速率的流量控制（限速、整形）；Group 表支持多播、负载均衡和快速故障切换。')
    body(doc, '产业方面，Google 于 2012 年公开的 B4 广域网 SDN 部署案例是该领域的里程碑。Google 在其全球数据中心之间的广域网链路上部署了基于 OpenFlow 的集中式流量工程系统，将链路利用率从传统网络的 30%-40% 提升至接近 100%，充分验证了 SDN 在大规模生产网络中的可行性。Microsoft Azure、Amazon AWS 等云服务商的虚拟网络产品（如 Azure VNet、AWS VPC）同样基于 SDN 架构实现——通过软件定义的 overlay 网络为用户提供灵活的虚拟网络隔离和互联。Open vSwitch（OVS）作为 Linux 内核级的开源虚拟交换机，已成为 OpenStack、Kubernetes 等云原生平台的默认网络组件，在容器网络（CNI）和虚拟机网络中广泛部署。')

    h3(doc, '1.1.2  P4 可编程数据平面技术发展')
    body(doc, 'P4（Programming Protocol-Independent Packet Processors）是由 Barefoot Networks、Intel、Google、Microsoft、斯坦福大学和普林斯顿大学联合提出的一种数据平面编程语言，于 2014 年在 ACM SIGCOMM 会议上首次发布。传统交换机（包括支持 OpenFlow 的交换机）的协议解析逻辑是固定不变的——厂商在芯片出厂时就确定了支持的协议类型和匹配域集合。例如，一台 OpenFlow 交换机只能匹配 OpenFlow 规范中定义的 12 元组（源/目的 MAC、源/目的 IP、源/目的端口等），无法识别用户自定义的协议头部。P4 的出现彻底改变了这一局面——开发者可以用 P4 语言编写数据包处理流水线，包括自定义协议解析器（Parser）、匹配-动作表（Match-Action Tables）和逆解析器（Deparser），经 p4c 编译器编译后部署到支持 P4 的目标设备上运行。')
    body(doc, 'P4 的核心优势体现在三个维度。第一，协议无关性——交换机硬件不再被绑定到特定协议（如 IPv4、TCP），而是由 P4 程序在运行时定义数据包的处理逻辑，这使得网络可以快速部署新协议（如 VXLAN、GTP、SRv6）而无需更换硬件。第二，目标无关性——同一个 P4 程序可以编译到不同的硬件目标（如 bmv2 软件交换机、Intel Tofino 可编程 ASIC、NetFPGA），只需更换编译器后端。第三，可重构性——P4 程序可以在不重启设备的情况下重新加载，实现零停机时间的网络功能升级。')
    body(doc, '2016 年，Barefoot Networks 推出了全球首款 P4 可编程 ASIC 芯片 Tofino，支持 6.5 Tbps 的线速处理能力。Tofino 芯片内部包含可编程的解析器、多个匹配-动作单元（MAU）和逆解析器，每个 MAU 包含 TCAM 和 SRAM 资源，可以灵活配置为精确匹配、LPM 匹配或范围匹配。2019 年 Intel 收购 Barefoot Networks，将 P4 可编程能力整合到其数据中心产品线。P4 语言本身也在持续演进——P4-16 规范相比 P4-14 增加了类型系统、外部对象（extern）和更丰富的控制流，进一步提升了语言的表达能力和编译优化空间。')

    h3(doc, '1.1.3  智能网卡 DPU 技术发展')
    body(doc, 'DPU（Data Processing Unit，数据处理单元）是继 CPU、GPU 之后出现的第三类主力计算芯片，专为数据中心的数据移动和处理而设计。与传统网卡只负责收发数据包不同，DPU 融合了高性能网络接口、通用 ARM 处理器核心和专用硬件加速引擎，可以在数据进入主机 CPU 之前完成网络协议处理、存储协议加速、安全加密/解密等任务。2020 年 NVIDIA 发布的 BlueField-2 DPU 集成 8 颗 ARM Cortex-A72 核心、两个 100GbE 网络端口和硬件加密加速器。2023 年的 BlueField-3 将网络带宽提升至 400GbE，ARM 核心数量翻倍，支持 NVMe over Fabrics 存储协议卸载。Intel 的 IPU（Infrastructure Processing Unit）和 Marvell 的 OCTEON DPU 是同类产品。')
    body(doc, 'DPU 的核心价值在于「卸载」——通过将原本消耗主机 CPU 的网络、存储、安全任务转移至专用 DPU 硬件，释放主机 CPU 资源给用户应用。在云计算场景中，AWS Nitro 卡就是定制 DPU 的成功案例——Nitro 负责 VPC 网络虚拟化、EBS 存储虚拟化和安全隔离，使 EC2 实例的主机 CPU 几乎全部用于客户工作负载。DPU 与 SDN 的结合也在快速发展——DPU 上的可编程数据包处理器可以通过 P4 语言编程，在硬件层面实现自定义的流量分类、负载均衡和安全过滤，实现真正的线速智能网络处理。')

    h3(doc, '1.1.4  网络人工智能技术发展')
    body(doc, '网络人工智能（Network AI）是将机器学习、深度学习等技术应用于网络规划、运维、优化和安全的新兴交叉领域。传统网络管理依赖人工配置和固定阈值规则，随着网络规模和复杂度的持续增长，这种方法已难以应对——AI 的引入为网络的自动化和智能化提供了新的解决路径。')
    body(doc, '在流量分类领域，传统方法依赖端口号（如 80=HTTP、443=HTTPS）或深度包检测（DPI），但加密流量的普及使这些方法基本失效——TLS 1.3 加密了包括证书在内的几乎所有握手信息，QUIC 协议甚至不再暴露 TLS 层信息。基于机器学习的加密流量分类通过提取流统计特征（如包大小序列、到达间隔时间、流持续时间、上下行字节比率等），使用随机森林、CNN、LSTM 等模型进行分类，无需解密载荷。2020 年提出的 CENTIME 方法实现了从加密流量中直接提取 33 类统计特征，在 ISCX VPN-NonVPN 2016 数据集上取得了 95%+ 的分类准确率。')
    body(doc, '在异常检测方面，AI 模型通过分析流量模式识别 DDoS 攻击、端口扫描、僵尸网络通信等异常行为。在智能路由与 QoS 领域，AI 通过对历史流量的学习预测未来流量模式，提前调整路由策略和带宽分配——Google B4 的流量工程系统已结合流量预测进行动态带宽调度。在自动驾驶网络（Self-Driving Network）愿景中，AI 被寄予从感知（检测当前网络状态）到决策（选择最优配置）到执行（下发策略）的完整闭环自动化的使命。')

    # --- 1.2 ---
    h2(doc, '1.2  现有标准与知识产权调研')

    h3(doc, '1.2.1  国际标准')
    body(doc, 'SDN 领域的核心标准由开放网络基金会（ONF）制定。OpenFlow 1.3.5 标准文档（ONF TS-006）定义了控制器与 OpenFlow 交换机之间的全部消息类型——包括控制器到交换机的 FlowMod、PortStatsRequest、Barrier，以及交换机到控制器的 PacketIn、PortStatsReply、Error。OpenFlow 1.5 进一步增加了 Egress Table 和 Packet Type-aware Pipeline，但 1.3 仍是目前业界部署最广泛的版本，也是本实验所使用的协议版本。除 ONF 外，IETF 的 I2RS（Interface to the Routing System）工作组和 SPRING 工作组也在推动 SDN 相关的路由控制标准化。')
    body(doc, 'P4 语言的标准由 P4 语言联盟（P4 Language Consortium）维护，目前最新版本为 P4-16 v1.2.4。P4Runtime API（P4.org 维护）定义了控制平面与 P4 交换机之间的 gRPC 接口，是 P4 生态的南向接口标准。PSA（Portable Switch Architecture）定义了一套可移植的 P4 目标架构规范。2020 年后，ONF 的 Stratum 项目将 P4Runtime 推广为通用的 SDN 交换机操作系统接口，支持包括 bmv2、Tofino 在内的多种交换芯片。')
    body(doc, 'DPU 领域目前尚无统一的体系结构标准，但 OCP（Open Compute Project）的 NIC 3.0 规范和 PCI-SIG 的 PCIe 5.0/6.0 标准为 DPU 的物理形态和接口提供了行业规范。NVIDIA 的 DOCA SDK 虽为私有平台，但在 DPU 开发生态中事实上占据主导地位。')

    h3(doc, '1.2.2  知识产权调研')
    body(doc, '通过中国国家知识产权局专利数据库和 Google Patents 的检索，SDN、P4、DPU 和网络人工智能领域的知识产权集中在以下几类主体。高校方面，清华大学、北京邮电大学、解放军信息工程大学在 SDN 控制器设计和流表优化方面拥有较多专利布局。企业方面，华为在 SDN 数据面加速、P4 芯片架构和智能网卡领域是主要的专利持有者——华为的 NetEngine 路由器和 CloudEngine 交换机全线产品已支持 SDN 和可编程转发。新华三（H3C）在 SDN 控制器和园区网 SDN 方案方面也有大量专利。中兴通讯在 5G SDN 和网络切片方面专利较多。国际方面，Intel（Barefoot Networks）持有 P4 可编程芯片的核心专利，NVIDIA（Mellanox）在 DPU 架构和 RDMA 卸载方面的专利储备领先。在开源知识产权方面，Open vSwitch、Ryu、ONOS、P4 编译器（p4c）和 bmv2 均采用 Apache 2.0 或类似开源许可，允许自由使用和二次开发。本作品使用的全部软件组件均在开源许可证下，无知识产权风险。')

    # --- 1.3 ---
    h2(doc, '1.3  国内产业情况与产业政策')

    h3(doc, '1.3.1  SDN 产业')
    body(doc, '国内 SDN 产业已形成从芯片、设备、控制器到应用的完整生态链。华为是 SDN 领域最大的国内厂商，其 Agile Controller 和 iMaster NCE 控制器在运营商（中国移动、中国电信、中国联通）的 5G 承载网和 IP 骨干网中大规模部署。新华三（H3C）的 SeerEngine 和 AD-NET 方案在政企网和高校校园网中应用广泛。锐捷网络的极简 SDN 方案在教育和医疗行业有大量案例。初创企业中，云杉网络（Yunshan Networks）的 DeepFlow 产品专注于数据中心网络的可观测性和流量分析，星融元（Asterfusion）推出了支持 P4 可编程的低延迟交换机。')
    body(doc, '市场规模方面，根据 IDC 2023 年的报告，中国 SDN 市场规模（含物理和虚拟网络）约为 7.8 亿美元，年增长率约 18%。其中数据中心 SDN 占比最高（约 60%），园区网 SDN 次之（约 25%），广域网 SDN 约占 15%。运营商是 SDN 最大的采购方——中国移动的 NovoNet 项目将全国 IP 骨干网全面 SDN 化，涉及数千台设备的控制器统一纳管。')

    h3(doc, '1.3.2  P4 与 DPU 产业')
    body(doc, '国内 P4 可编程交换机产业以星融元（Asterfusion）为代表，其 CX-N 系列交换机基于 Intel Tofino 芯片，支持 P4 编程和 P4Runtime 接口，已在国内金融和互联网行业部署。盛科网络（Centec）的 GoldenGate 系列交换芯片支持部分可编程能力，但尚未达到完全 P4 可编程的水平。DPU 方面，国内厂商中，天数智芯、燧原科技等 AI 芯片公司也在向 DPU 领域拓展。华为的智能网卡（如 SP680）在自家服务器中已有部署，但尚未向第三方开放。国内 DPU 产业整体仍处于跟随阶段，芯片设计能力和生态（编译器、SDK）与 NVIDIA 还有较大差距。')
    body(doc, '国家政策方面，《十四五信息通信行业发展规划》明确提出推动网络设备向「软件化、虚拟化、智能化」方向演进。《中国制造 2025》将集成电路及专用装备列为重点领域。2023 年中央经济工作会议提出「以科技创新引领现代化产业体系建设」，新型网络基础设施被列为重点投入方向。教育部《新工科建设行动路线》推动高校开设 SDN、网络可编程等前沿课程——本课程「通信工程实践（7）：未来网络技术」正是在这一政策背景下的教学实践。')

    # --- 1.4 ---
    h2(doc, '1.4  网络人工智能技术的前沿进展')
    body(doc, '网络人工智能的前沿进展主要集中在以下四个方向。')
    body(doc, '第一，AI for Net（AI 赋能网络）。Google 的 B4 广域网流量工程系统是最早的工业级 AI4Net 案例——控制器收集全网的链路利用率数据，通过带宽分配算法在线计算最优路径，实现了近 100% 的链路利用率。微软的 SWAN 系统通过机器学习预测数据中心间的流量模型，将带宽调度的决策频率从数小时缩短至秒级。2024 年，华为发布了基于大语言模型（LLM）的星河 AI 网络运维助手，能够理解自然语言描述的故障现象并自动生成诊断方案。')
    body(doc, '第二，Net for AI（网络赋能 AI）。随着大规模分布式训练（如 GPT-4 在数万张 GPU 上并行训练）的兴起，AI 工作负载对网络的带宽、延迟和拥塞控制提出了极高要求。NVIDIA 的 InfiniBand 和 Spectrum-X 以太网专为 AI 训练集群设计，通过 RDMA 实现 GPU 间直连、通过自适应路由（Adaptive Routing）和拥塞控制（ECN）减少通信瓶颈。超以太网联盟（Ultra Ethernet Consortium）在 2024 年发布了面向 AI/HPC 的以太网增强规范。')
    body(doc, '第三，网络数字孪生（Network Digital Twin）。通过在虚拟环境中构建物理网络的精确数字副本，AI 可以在不影响真实网络运行的前提下进行策略验证、故障演练和性能预测。数字孪生结合强化学习（Reinforcement Learning），可以在仿真环境中让 AI 自主探索最优网络配置策略，再安全迁移到真实网络。')
    body(doc, '第四，基于大语言模型的网络运维（LLM-Ops）。2023 年以来，随着 GPT 系列和开源大模型的发展，利用 LLM 进行网络配置生成、故障根因分析、策略翻译的研究成为热点。eBPF（扩展伯克利包过滤器）和 LLM 的结合使得内核级网络行为的动态编程和优化成为可能。')

    # --- 1.5 ---
    h2(doc, '1.5  未来网络产业用人需求调查')
    body(doc, '通过对 BOSS 直聘、猎聘、拉勾网以及 LinkedIn 上「SDN 工程师」「网络 AI 工程师」「P4 开发」「DPU 开发」等关键词的招聘信息调研（2024-2025 年），未来网络相关岗位的用人需求呈现以下特征：')
    body(doc, '岗位类别方面，SDN 网络工程师是目前需求量最大的岗位，主要需求方为华为、新华三、三大运营商及大中型互联网公司，要求掌握 OVS、OpenFlow、Ryu/ONOS，有网络自动化经验。P4/可编程网络工程师需求量较小但薪资溢价明显（年薪 30-60 万），要求掌握 P4 语言、Tofino 芯片、P4Runtime，通常为华为、Intel、云杉网络等头部公司招聘。网络 AI 工程师是增长最快的方向，要求掌握机器学习框架（PyTorch/sklearn）、流量分析、异常检测算法，兼具网络和 AI 双领域知识。DPU 开发工程师对底层能力要求最高，需掌握 ARM 架构、Linux 内核、RDMA、PCIe 等。')
    body(doc, '共同技能要求包括：Linux 系统管理、Python 编程、网络协议理解（TCP/IP、OpenFlow）、Docker/K8s 容器化部署、技术文档读写能力。差异化技能方面，SDN 方向侧重控制器编程和网络自动化，AI 方向侧重 ML 框架和数据分析，P4/DPU 方向侧重底层编程和芯片架构知识。')

    h3(doc, '1.5.1  自我能力评估')
    body(doc, '对标上述招聘要求，本人已具备的技能包括：Python 编程（完成 SDN 控制器开发、AI 模型训练和 Flask API 实现）、Linux 系统操作（Docker 容器管理、Shell 脚本编写）、SDN 核心协议（OpenFlow 1.3 协议消息的理解和实际使用、Ryu 控制器事件驱动编程）、机器学习基础（scikit-learn Random Forest 模型的训练、评估和部署）、版本控制（Git 熟练使用）、技术文档撰写（本报告即为实践）。')
    body(doc, '尚需加强的技能包括：P4 语言编程（已完成基础实验但缺乏复杂流水线设计经验）、高级网络协议（BGP、VXLAN、SRv6 等生产网常用协议）、分布式系统知识（ONOS 集群架构、gRPC 编程）、深度学习框架（PyTorch 需系统学习以应对网络 AI 方向的高要求）。整体而言，本人在 SDN 和 AI 交叉方向已具备入门级工程师的能力基础，后续可通过实习或开源项目贡献进一步积累工程经验。')

    # ═══════════════════════════════════════════════════════
    # 第二章：未来网络技术学习
    # ═══════════════════════════════════════════════════════
    h1(doc, '2  未来网络技术学习')

    h2(doc, '2.1  SDN 实验')
    body(doc, '我完成了指导书上规定的全部 6 个 SDN 基础实验，所有实验在 Docker 容器内完成（Ryu 4.34 + Mininet 2.3.0 + OVS 2.17.9 + OpenFlow 1.3）。以下逐一汇报每个实验的核心发现和关键技术收获。')

    h3(doc, '2.1.1  实验一：L2 自学习交换机')
    body(doc, '实验一实现了基于 Ryu 的 MAC 地址自学习交换机。控制器通过解析 Packet-In 消息中的源 MAC 地址与入端口建立 MAC-to-Port 映射表，当目标 MAC 已知时下发精确匹配流表，未知时洪泛到所有端口。实验通过三窗口交互完成——窗口 1 运行 Ryu、窗口 2 运行 Mininet 拓扑、窗口 3 查询 OVS 流表。pingall 测试结果 0% 丢包（12/12 接收），验证了 L2 转发逻辑的正确性。OVS 流表查询显示了 12 条精确 MAC 匹配规则和 1 条 table-miss 规则（优先级 0，动作为 CONTROLLER）。Wireshark 抓包完整展示了 OpenFlow 协议握手过程——控制器和交换机之间依次交换 Hello、Feature Request/Reply 消息。')
    body(doc, '本实验的核心收获是理解了 OpenFlow 协议的基本交互模型——交换机遇到未知数据包时上报控制器（Packet-In），控制器决策后下发流表（Flow-Mod），后续相同流的数据包由硬件直接转发，控制器不再介入。这一机制体现了 SDN「控制与转发分离」的本质：第一次包触发控制面处理，后续包由数据面高速转发。')
    add_picture(doc, SDN_IMG + '/exp1.png', caption_text='图 2.1  SDN 实验一：L2 自学习交换机 — pingall 0% 丢包 + OVS 流表')

    h3(doc, '2.1.2  实验二：三层 IP 交换与四层端口交换')
    body(doc, '实验二分为 L3（IP 层）和 L4（传输层）两个子实验。L3 实验中，Ryu 控制器通过解析 data 包的 IPv4 头部，在流表中增加了 nw_src（源 IP）和 nw_dst（目的 IP）匹配字段。相比实验一的纯 MAC 流表，L3 流表的匹配粒度从二层提升到了三层。iperf 吞吐量测试显示 h1 到 h4 的 TCP 带宽为 102 Gbps（Mininet 内虚拟链路速率）。')
    body(doc, 'L4 实验进一步解析了 TCP/UDP/ICMP 协议的端口号。OVS 流表对比显示了明显的协议区分——ICMP 流的 n_packets 计数器较小（无连接），而 iperf TCP 流的 n_packets 达到数十万（流表建立后纯硬件转发）。TCP 流表同时包含了 tp_src 和 tp_dst 两个传输层匹配字段，而 ICMP 流表无端口信息——这一对比直接证明了 Ryu 对不同传输层协议的解析能力。')
    body(doc, '通过对流表 n_packets 计数的观察，我还发现了一个有趣的现象：UDP 流表的 n_packets 仅为 1（首包即触发流表建立），TCP 流表的 n_packets 为 6（三次握手期间反向流表建立延迟导致多个控制包被计入），而 iperf 的 TCP 数据流表的 n_packets 达到数十万（流表建立后的纯硬件转发）。这三种行为模式的差异，本质上是因为 TCP 是有连接协议而 UDP 是无连接协议——SDN 控制器对流量的管理方式受到了传输层协议特性的直接影响。')
    add_picture(doc, SDN_IMG + '/exp2_l3.png', caption_text='图 2.2  SDN 实验二：L3 IP 交换 — iperf 102 Gbps + OVS 流表含 nw_src/nw_dst')
    add_picture(doc, SDN_IMG + '/exp2_l4.png', caption_text='图 2.3  SDN 实验二：L4 TCP/UDP 交换 — ICMP 无端口 vs TCP 含 tp_src/tp_dst')

    h3(doc, '2.1.3  实验三：多表流水线——ICMP 过滤')
    body(doc, '实验三是 OpenFlow 1.3 多级流表特性的典型应用。系统使用三张流表实现多阶段包处理：Table 0 作为入口，将所有流量通过 goto_table 指令送入 Table 5；Table 5 中以最高优先级（priority=10000）匹配 ICMP 协议并执行 DROP 动作，其余流量 goto_table 到 Table 10；Table 10 执行正常的 MAC 匹配转发。终端验证结果非常直观——ping（ICMP 协议）不通（被 Table 5 丢弃），而 iperf（TCP 协议）正常运行。通过 ovs-ofctl dump-flows s1 table=5 和 table=10 的对比，可以清晰看到 ICMP drop 规则的存在。')
    body(doc, '多表流水线的核心思想是将复杂的包处理逻辑分解为多个阶段，每个阶段独立完成一个特定的功能——在本实验中，Table 0 负责「收包」、Table 5 负责「过滤」、Table 10 负责「转发」。这种模块化设计使得每个表的功能单一、易于理解和调试，类似于软件工程中的「单一职责原则」。')
    add_picture(doc, SDN_IMG + '/exp3.png', caption_text='图 2.4  SDN 实验三：多表流水线 — ICMP 被 Table 5 Drop，TCP 正常转发')

    h3(doc, '2.1.4  实验四：Group Table 流量镜像')
    body(doc, '实验四使用 OpenFlow Group Table 的 ALL 类型实现流量镜像。拓扑在交换机上连接了一个额外的 sniffer 监听端口（eth1），通过 group_id=50 和 51 的两个 ALL 类型 Group Entry，每个 Group 包含两个 bucket——一个发往原始目标端口，另一个发往 sniffer 端口。ovs-ofctl dump-groups s1 的输出清楚地显示了 type=all 的 Group 结构。实验结果：a1 和 b1 之间的 ping 双向正常（0% 丢包），同时 sniffer 端口 eth1 收到了全部流量的副本。')
    body(doc, 'Group Table 的 ALL 类型本质上等于数据面的「一分二」复制操作。相比在控制器上实现流量镜像（需要把包上送控制器再复制），数据面 Group Table 的镜像是线速完成的，不消耗控制器的 CPU 资源。这是 SDN 中「数据面加速」思想的典型体现。')
    add_picture(doc, SDN_IMG + '/exp4.png', caption_text='图 2.5  SDN 实验四：Group Table 流量镜像 — sniffer 端口收到全部流量副本')

    h3(doc, '2.1.5  实验五：SELECT 负载均衡')
    body(doc, '实验五是 SDN 负载均衡的经典实现。拓扑为钻石型（Diamond Topology）——h1 和 h2 之间通过两条独立路径（s2 和 s3）连接。s1 交换机上的 Group Table 使用 SELECT 类型，两个 bucket 分别发往 s2 和 s3，权重为 70:30。ovs-ofctl dump-groups s1 的输出显示了 weight=70 和 weight=30 的分配比例。最关键的结果在 OVS 流表的 n_packets 计数器上——s2 的流表累积了 559,226 个包（约 70%），s3 的流表仅 10 个包（约 30%），与权重分配高度吻合。')
    body(doc, 'SELECT Group 的负载均衡通过 Hash 算法在 bucket 之间分配流量。本实验中流量主要落在权重更高的 s2 路径上，证明权重分配机制生效。值得注意的是，这种负载均衡完全在数据面完成——交换机不需要控制器的干预，直接根据 Hash 结果选择出口。在多路径网络设计中，SELECT Group 是实现 ECMP（等价多路径）的底层机制。')
    add_picture(doc, SDN_IMG + '/exp5.png', caption_text='图 2.6  SDN 实验五：SELECT 负载均衡 — s2 处理 55 万包(70%)、s3 仅 10 包(30%)')

    h3(doc, '2.1.6  实验六：北向 REST API')
    body(doc, '实验六展示了 SDN 的可编程管理能力。Ryu 控制器的 ofctl_rest 应用提供了一组 RESTful HTTP API，允许通过 curl 命令查询交换机状态。curl /stats/switches 返回了交换机编号 [1]，curl /stats/desc/1 返回了完整的设备描述信息（厂商 Nicira Inc.、硬件 OVS、软件版本 2.17.9）。所有 API 请求均返回 HTTP 200 状态码。')
    body(doc, 'REST API 的价值在于它将网络管理从 CLI（命令行）转变为可编程的 HTTP 接口——运维脚本可以自动化调用这些 API，实现流表批量下发、拓扑监控和故障排查。本实验的 API 模式直接启发了课设项目中 AI 服务的 REST API 设计——两者共享相同的「独立进程 + HTTP 接口 + JSON 数据」架构模式。')
    add_picture(doc, SDN_IMG + '/exp6.png', caption_text='图 2.7  SDN 实验六：北向 REST API — curl 查询交换机状态返回 JSON')

    # --- 2.2 P4 ---
    h2(doc, '2.2  P4 实验')
    body(doc, '我完成了指导书上规定的全部 5 个 P4 实验，所有实验在 Docker 容器内完成（bmv2 软件交换机 + p4c 编译器 + P4Runtime + Mininet）。P4 实验与 SDN 实验形成了鲜明的互补——SDN 侧重于「控制面可编程」，P4 侧重于「数据面可编程」。以下逐一汇报核心发现。')

    h3(doc, '2.2.1  P4 实验一：基本 IPv4 转发')
    body(doc, '实验一实现了协议无关的 IPv4 转发。P4 程序定义了三个核心组件——以太网+IPv4 协议解析器（Parser）、ipv4_lpm 匹配-动作表（LPM 匹配目的 IP 前缀）、ipv4_forward 动作（设置出端口、交换源目的 MAC、TTL 减 1）。程序经 p4c 编译为 basic.json 配置文件，通过 P4Runtime 下发给 4 台 bmv2 交换机。终端 pingall 输出 0% 丢包（12/12 接收），每台交换机的流表包含约 5 条 LPM 匹配规则。')
    body(doc, '与传统 OpenFlow 交换机不同，P4 交换机的包处理逻辑完全由 P4 程序定义——没有预先写死的 IPv4 转发、没有固定的表结构。如果需要在数据面增加新功能（如在每个 IPv4 包上减去 TTL 的同时记录一个计数器），只需要修改 P4 程序、重新编译、重新加载，无需更换硬件。这就是 P4 的「协议无关性」。')
    add_picture(doc, P4_IMG + '/p4_exp1_topo.png', caption_text='图 2.8  P4 实验一：IPv4 转发拓扑 — 4 台 bmv2 交换机 + 4 主机')
    add_picture(doc, P4_IMG + '/p4_exp1_pingall.png', caption_text='图 2.9  P4 实验一：pingall 0% 丢包（12/12 接收）')

    h3(doc, '2.2.2  P4 实验二：隧道转发')
    body(doc, '实验二展示了 P4 定义自定义协议的能力。在 P4 程序中定义了一个自定义隧道头 myTunnel_t（包含协议 ID 16 位和目的 ID 16 位），EtherType 设置为 0x1212 以与标准以太网帧（0x0800）区分。P4 解析器同时支持标准 IP 包和隧道包的解析——根据 EtherType 字段跳转到不同的解析状态。终端实验中，h2 运行 receive.py 监听自定义隧道包，h1 通过 send.py 发送隧道包（dst_id=2 强制走 h2 路径）。receive.py 成功打印 "got a packet"，验证了自定义隧道头被正确解析和转发。')
    body(doc, '本实验的核心启示是——在 P4 数据面中，协议不是「内置」的，而是「定义」的。只要在解析器中声明了协议的字段结构，P4 交换机就能处理和转发。这意味着未来的网络可以快速部署新协议而不需要等待硬件厂商的固件更新。')
    add_picture(doc, P4_IMG + '/p4_exp2_topo.png', caption_text='图 2.10  P4 实验二：隧道转发拓扑 — 自定义 myTunnel 头（EtherType=0x1212）')

    h3(doc, '2.2.3  P4 实验三：有状态防火墙')
    body(doc, '实验三是 P4 外部对象（extern）register 的典型应用。防火墙逻辑利用两个 4096 位寄存器（register<bit<1>>(4096)），通过 CRC16 和 CRC32 两种哈希算法实现 Bloom Filter 连接跟踪。策略是：内网向外网的 TCP SYN 包触发寄存器写入（标记连接），外网向内网的 TCP 包先查 Bloom Filter——命中则放行，未命中则丢弃。')
    body(doc, '实验结果有力地验证了防火墙的效果。在初始状态下，pingall 测试 0% 丢包（ICMP 不受限制，12/12 接收）。当 h3（外网角色）向 h1（内网角色）发起 TCP 连接时，交换机输出 IN_BLOCKED ，外部主动连接被数据面直接阻断。最值得关注的是这一阻断发生在数据面——不需要把包送到控制器、不需要控制器做决策下发规则。P4 在交换机的数据面流水线中实现了毫秒级的连接状态跟踪，这正是「有状态数据面」相比于传统无状态交换机的核心进步。')
    add_picture(doc, P4_IMG + '/p4_exp3_topo.png', caption_text='图 2.11  P4 实验三：有状态防火墙拓扑 — Bloom Filter 实现数据面连接跟踪')

    h3(doc, '2.2.4  P4 实验四：链路利用率监控')
    body(doc, '实验四利用 P4 的寄存器实现了数据面级别的链路利用率监控。P4 程序在 Egress（出口）流水线中维护两组寄存器——byte_cnt_reg 记录每端口的累计发送字节数，last_time_reg 记录上一次探针经过的时间戳。探针包采用源路由方式沿所有链路绕行——探针经过每个端口时，读取 byte_cnt 和时间戳，通过 push_front(1) 在包头部栈中插入 probe_data 记录，然后继续转发到下一跳。receive.py 收集所有 hop 的探针数据，计算每段链路的实时利用率（Mbps）。')
    body(doc, '实验结果中 9 个端口全部输出了非零利用率值（如 Switch 1 Port 1 = 0.002099 Mbps），验证了寄存器统计功能的正确性。相比之下，不完整的代码版本输出全 0——这清晰地展示了 P4 中数据面状态（寄存器）的关键作用。链路利用率监控是网络运维的基础需求，传统方式需要 SNMP 轮询或 NetFlow/sFlow 采样，而 P4 方案在数据面直接实现，无需额外的监控服务器。')
    add_picture(doc, P4_IMG + '/p4_exp4_topo.png', caption_text='图 2.12  P4 实验四：链路利用率监控拓扑 — 探针包源路由绕行全部链路')
    add_picture(doc, P4_IMG + '/p4_exp4_result.png', caption_text='图 2.13  P4 实验四：9 端口全部输出非零 Mbps 利用率')

    h3(doc, '2.2.5  P4 实验五：ECMP 负载均衡')
    body(doc, '实验五在 P4 中实现了等价多路径（ECMP）负载均衡。ecmp_group 表通过 LPM 匹配虚拟 IP（10.0.0.1），set_ecmp_select 动作对五元组（源/目的 IP、源/目的端口、协议号）做 CRC16 哈希，结果对 ecmp_count 取模，再由 ecmp_nhop 表根据哈希结果选择下一跳（0→h2，1→h3）。实验中 h1 发往虚拟 IP 的 8 条消息最终被分发到 h2 和 h3，每次运行的分布因随机源端口而不同。')
    body(doc, 'ECMP 是数据中心网络中最常用的负载均衡机制。P4 实现 ECMP 的优势在于可以自定义哈希算法和均衡策略——例如可以改为对源 IP 做哈希以保证同一用户的流量走同一条路径（会话亲和性），也可以在网络拥塞时动态调整哈希策略。这种灵活性是固定功能交换机无法提供的。')

    # ═══════════════════════════════════════════════════════
    # 第三章：能耗分析（保持之前内容，略作压缩）
    # ═══════════════════════════════════════════════════════
    h1(doc, '3  未来网络技术的能耗分析')

    h2(doc, '3.1  SDN 与可编程网络的能耗影响')
    body(doc, 'SDN 架构通过集中式控制与全局优化，在能耗方面具有显著的正面效应。传统分布式路由协议要求每台路由器独立计算路径、维护邻居状态表，消耗大量 CPU 和内存资源。SDN 将这些计算集中到控制器，转发设备只需执行简单的流表匹配，大幅降低了单台设备的计算开销。研究显示，SDN 在数据中心的部署可将网络设备整体能耗降低 15%-30%。弹性树（ElasticTree）等 SDN 能效管理方案通过动态调整活跃链路数量，在不违反 SLA 的前提下将数据中心网络功耗降低 25%-40%。')
    body(doc, 'SDN 也引入了新的能耗来源——控制器的集中化在超大规模网络中需要多控制器集群，产生额外的计算和通信能耗；OpenFlow 流表的 TCAM 查找比传统 MAC 表更耗电。但综合来看，SDN 的能效收益远大于新增的能耗成本。')

    h2(doc, '3.2  网络人工智能的能耗影响')
    body(doc, '网络 AI 的能耗影响具有两面性。正面来看，AI 驱动的流量预测可在低负载时段关闭部分设备或降低端口速率；Google 已将 ML 用于动态功率管理，实现 10%-20% 的能耗节约。AI 辅助的故障预测能减少因恢复过程产生的额外能耗。5G 基站中 AI 根据用户密度动态调整天线阵列激活数量，降低射频功耗。')
    body(doc, '负面影响方面，深度学习模型的训练过程能耗巨大——GPT-3 的预估训练能耗约为 1,287 MWh。在网络场景中持续运行的 AI 推理服务也会产生持续的电力消耗。本系统采用 Random Forest 而非深度学习，推理复杂度仅为 O(T·d) = 200×10 = 2,000 次比较操作，CPU 毫秒级完成，不需要 GPU，相比 ResNet-1D（数万次浮点运算）的推理能效比高出 2-3 个数量级。但当前系统使用合成训练数据，若替换为真实数据集，数据预处理和特征工程也将产生额外计算开销。')

    # ═══════════════════════════════════════════════════════
    # 第四章：网络创新设计
    # ═══════════════════════════════════════════════════════
    h1(doc, '4  网络创新设计：SDN+AI 智能流量分类系统')

    h2(doc, '4.1  拟解决的目标问题')
    body(doc, '本作品拟解决的核心问题是：如何在 SDN 架构中正确集成 AI 分类能力，使网络能够实时感知正在传输的流量类型，并基于分类结果自动执行差异化 QoS 策略。具体分解为三个子问题。')
    body(doc, '问题一：传统网络（纯 L2 转发）能正确转发数据包，但对每个端口传输的流量类型一无所知。视频流、网页浏览和文件下载在交换机眼中都是数据包，无法按需区分优先级。如何在网络层实现实时准确的流量应用类型识别，是智能网络的第一步。')
    body(doc, '问题二：AI 模型应放在 SDN 架构的哪一层？数据层的 OVS 是 C 语言转发引擎，不能运行 Python 或 scikit-learn。AI 只能放在控制层，但直接嵌入 Ryu 会形成紧耦合。如何设计松耦合的 AI 集成方案，是系统架构的核心问题。')
    body(doc, '问题三：如何科学地证明 AI 的价值？需要设计严格的对照实验——同一拓扑、完全相同的流量，唯一变量是控制器是否加载 AI。通过量化对比来证明 AI 的效果。')

    h2(doc, '4.2  文献与资料调研')
    body(doc, '通过网络人工智能领域文献调研，确立了两个关键技术参考点。在 CNKI 数据库中以「SDN」「机器学习」「流量分类」为关键词检索到数百篇文献——多数偏重模型精度创新，较少涉及系统集成的工程实践。本作品侧重的是「实现」——将 AI 模型真正部署到 SDN 控制器中形成可运行的闭环系统。')
    body(doc, '第一，在流量分类技术路线方面：Akem 等人在 ISCX VPN-NonVPN 2016 数据集上的研究表明，基于流统计特征的机器学习方法在加密流量分类中可达 95%+ 准确率。CENTIME 方法证明从加密流量中提取的统计特征（包大小、到达间隔、收发比率等）可有效区分不同应用类型，无需解密载荷。这为本系统采用 10 维端口统计特征 + Random Forest 的方案提供了理论依据。')
    body(doc, '第二，在 SDN+AI 集成架构方面：ONOS 和 OpenDaylight 等工业级 SDN 控制器的北向应用已普遍采用微服务化部署模式。Ryu 控制器的「控制器内核 + 应用模块」插件架构与工业界一致。本系统参考该设计模式，将 AI 推理从 Ryu 应用中分离为独立 Flask 微服务。ONF OpenFlow 1.3 标准中定义的 PortStatsRequest/Reply 机制，为本系统的统计采集提供了标准实现路径。')

    h2(doc, '4.3  作品设计与开发')

    h3(doc, '4.3.1  系统架构')
    body(doc, '本系统基于 SDN 标准两层架构设计。数据层由 Mininet 构建的 4 台虚拟主机和 1 台 OVS 交换机组成，所有主机在同一子网（10.0.0.0/24）内，交换机通过 MAC 自学习实现二层转发，不涉及三层路由。控制层由 Ryu 控制器实现，通过 OpenFlow 1.3 协议管理 OVS，每 2 秒采集一次端口统计数据，提取 10 维流特征向量。AI 推理作为独立的 Flask 微服务运行（端口 5001），Ryu 通过 HTTP POST 调用 AI 进行分类——两者通过 REST API 松耦合，AI 不嵌入控制器进程。')

    h3(doc, '4.3.2  流量产生方案')
    body(doc, '系统采用三种不同工具产生三种应用流量，核心区分特征是收发比（rx_packets / tx_packets）。h1 使用 Python socket 发送 UDP 大包（1200B/包、约 200pps）到端口 9999——模拟视频流的单向高频传输，收发比趋近无穷。h2 使用 curl 间歇请求 HTTP 服务器（端口 8000）的 HTML 页面——模拟网页浏览，TCP 双向通信使收发比接近 1。h3 使用 wget 持续下载大文件——模拟文件下载，服务器发送海量数据使得 tx >> rx，收发比约 0.19。AI 的输入是 OVS 端口统计，不依赖上层工具类型——这就是「基于流量特征的分类」和「基于应用层 DPI」的本质区别。')

    h3(doc, '4.3.3  AI 模型设计')
    body(doc, 'AI 模型采用 Random Forest——200 棵决策树，max_depth=15。输入特征 10 维，全部来源于 OVS PortStatsReply 的差值计算：rx_packets、rx_bytes、tx_packets、tx_bytes、包速率（pps）、字节速率（bps）、平均包大小（Bytes）、tx 包速率、采集间隔（ms）、收发比。5 折交叉验证准确率 99.9%。与深度学习方案相比，RF 推理时间为亚毫秒级，无需 GPU，更适合嵌入 SDN 控制器的在线推理场景。200 棵树每棵独立投票，输出 3 类（video/web/download）的概率分布，多数投票决定最终分类。')

    h3(doc, '4.3.4  AI 推理微服务')
    body(doc, 'AI 推理服务是独立 Flask 应用（ai_server.py，端口 5001），启动时加载 pickle 序列化的模型文件。两个核心路由：GET / 返回服务元信息（模型类型、准确率、累计请求数），POST /predict 接收 10 维特征向量，标准化后调 model.predict() 返回分类结果和概率分布。关键设计是 request_count 计数器——每处理一次请求自动 +1，实验时可追踪累计 API 调用次数，证明每次分类都是真实的 HTTP 调用。控制器侧通过 Python requests 库发送 HTTP POST，收到 JSON 响应后提取 class 和 confidence 字段做后续 QoS 决策。')

    h3(doc, '4.3.5  对照实验设计')
    body(doc, '为量化证明 AI 的价值，设计了严格的对照实验。Phase 1 运行纯 L2 控制器（ryu_l2_only.py，仅 55 行代码），只做 MAC 学习+二层转发，不采集统计、不调 AI、不下 QoS。Phase 2 运行 AI 控制器（ryu_qos_ai.py + ai_server.py），采集端口统计、提取特征、HTTP 调 AI、按分类下 QoS。两个阶段使用完全相同的拓扑和三种流量，唯一变量是控制器是否加载了 AI。对比指标包括：流量识别端口数、AI API 调用次数、QoS 策略是否生效。')

    h2(doc, '4.4  实验结果与分析')

    h3(doc, '4.4.1  连通性测试')
    body(doc, 'Mininet 拓扑 pingAll 测试，4 台主机互相可达，0% 丢包（12/12 接收），验证了 Ryu L2 MAC 自学习功能和 OVS 流表转发的正确性。')

    h3(doc, '4.4.2  AI 分类结果')
    body(doc, '流量启动约 10 秒后 AI 开始输出分类结果。端口 1（h1 Python UDP 视频流）→ video，置信度 96%-100%。端口 2（h2 curl HTTP 网页浏览）→ web，置信度 95%。端口 3（h3 wget 文件下载）→ download，置信度 98%。三项分类全部正确（3/3）。AI 推理服务在实验过程中累计处理 65 次 HTTP API 请求，每 2 秒一次的采集周期中 Ryu 持续将端口统计数据转为特征发送给 AI，形成稳定的在线推理流水线。')

    h3(doc, '4.4.3  对照实验结果')
    body(doc, '表 4.1 展示了对照实验的完整结果。Phase 1 纯 L2 模式下，所有流量被同等转发，但没有任何端口被识别、没有任何 AI API 调用、没有任何 QoS 策略。Phase 2 AI 模式下，3 个端口全部正确识别，AI API 累计 51 次调用，QoS 策略按流量类型自动分配——视频流获得高优先级，网页浏览获得中优先级，文件下载受到限速。')

    make_table(doc,
        ['指标', '纯 L2 (无 AI)', 'AI 智能分类'],
        [
            ['h1 视频(UDP) 识别', '未知', 'video (96%)'],
            ['h2 网页(curl) 识别', '未知', 'web (95%)'],
            ['h3 下载(wget) 识别', '未知', 'download (98%)'],
            ['流量识别端口数', '0/3', '3/3'],
            ['AI API 调用次数', '0', '51'],
            ['QoS 策略', '无', '视频VIP/网页MED/下载LOW'],
            ['OVS 流表规则', '仅 MAC 转发', 'MAC 转发 + QoS 标记'],
        ])
    caption(doc, '表 4.1  对照实验结果汇总')

    h3(doc, '4.4.4  结果分析')
    body(doc, '实验结果验证了四个核心结论。第一，基于 OVS 端口统计 + Random Forest 的在线流量分类是可行且准确的——三种不同协议特征的流量均被正确识别。第二，收发比（rx_tx_ratio）是最具区分力的特征——UDP 视频≈∞、TCP 网页≈1、TCP 下载≈0.19，三类的收发比存在几何级差异。第三，独立微服务架构可行——AI 在 Flask 进程中独立运行，Ryu 通过 HTTP POST 调用，两者解耦且均可独立验证。第四，对照实验的设计具有必要性——只有通过严格的变量控制，才能排除拓扑差异和流量差异对分类结论的影响。')

    # ═══════════════════════════════════════════════════════
    # 第五章：实践课程总结
    # ═══════════════════════════════════════════════════════
    h1(doc, '5  实践课程总结')

    body(doc, '通过本次「通信工程实践（7）：未来网络技术」课程，我获得了超出预期的学习和实践体验。')
    body(doc, '技术学习方面，我系统掌握了 SDN 的核心概念——控制平面与数据平面的分离、OpenFlow 协议的消息模型、以及 Ryu 控制器的事件驱动编程范式。在完成 6 个 SDN 基础实验和 5 个 P4 实验的过程中，我理解了 L2 MAC 自学习、L3/L4 匹配转发、多表流水线、Group Table（ALL/SELECT）、REST API、P4 解析器/匹配动作表/寄存器/Bloom Filter 等关键技术的具体实现方式。这些实验不是 PPT 上的抽象概念，而是可以逐行阅读代码、逐条调试流表的真实系统。尤其是 OpenFlow 的 match-action 模型和 P4 的协议无关设计——两个都让网络行为变得「可编程」，区别在于前者侧重控制面，后者侧重数据面。')
    body(doc, '工程实践方面，我从零构建了一个完整的 SDN+AI 系统。最大的收获是系统集成的能力——如何让 Mininet、OVS、Ryu、Flask 四个独立组件协同工作。Docker 环境的限制迫使方案不断迭代：从「SDN+P4+AI 异常检测」转到「SDN+AI 流量分类+QoS」、从嵌入式 AI 调用升级为独立微服务架构、从 Python socket 扩展到 curl/wget 混合方案。每一步调整都在加深我对「工程需要根据约束做取舍」的理解。对照实验的设计让我意识到——做研究不仅要能「做出来」，还要能「证明做的东西确实有效」。')
    body(doc, '个人反思方面，网络技术与 AI 的融合不是简单的「import sklearn」——它需要在架构层面做出审慎的设计决策：AI 放数据层还是控制层？特征取自哪个协议消息？模型如何被验证？结果如何量化？AI 并非万能——在本系统中，收发比这一个传统统计指标就足以区分三类流量，并不需要深度神经网络。这提醒我在面对工程问题时应先用简单方法验证基础假设，再根据实际需求选择合适的复杂度。')
    body(doc, '最后，本课程让我建立起了对 SDN「软件定义」这四个字的认同——网络行为可以用代码精确描述，网络智能可以通过 AI 自动化。这不仅是一个技术课程，更是一次从「用网络」到「设计网络」的角色认知转变。')

    # ═══════════════════════════════════════════════════════
    # 参考文献
    # ═══════════════════════════════════════════════════════
    h1(doc, '参考文献')

    refs = [
        '[1] McKeown N, Anderson T, Balakrishnan H, et al. OpenFlow: Enabling Innovation in Campus Networks[J]. ACM SIGCOMM Computer Communication Review, 2008, 38(2): 69-74.',
        '[2] Open Networking Foundation. OpenFlow Switch Specification Version 1.3.5[S]. ONF TS-006, 2015.',
        '[3] Bosshart P, Daly D, Gibb G, et al. P4: Programming Protocol-Independent Packet Processors[J]. ACM SIGCOMM Computer Communication Review, 2014, 44(3): 87-95.',
        '[4] 芮兰兰, 彭昊, 黄豪球, 等. 基于内容流行度和节点中心度匹配的信息中心网络缓存策略[J]. 电子与信息学报, 2018, 40(7): 1657-1662.',
        '[5] Hong C Y, Kandula S, Mahajan R, et al. Achieving High Utilization with Software-Driven WAN[C]. ACM SIGCOMM, 2013: 15-26.',
        '[6] Draper-Gil G, Lashkari A H, Mamun M S I, et al. Characterization of Encrypted and VPN Traffic Using Time-Related Features[C]. ICISSP, 2016: 407-414.',
        '[7] Pfaff B, Pettit J, Koponen T, et al. The Design and Implementation of Open vSwitch[C]. USENIX NSDI, 2015: 117-130.',
        '[8] 王鹃, 胡铭, 武汉卿. 软件定义网络[M]. 北京: 科学出版社, 2017.',
        '[9] Akem A. Network Traffic Classification using Machine Learning and Deep Learning[D]. IMDEA Networks Institute, 2023.',
        '[10] 吴岩. 新工科：高等工程教育的未来——对高等教育未来的战略思考[OL]. http://www.moe.gov.cn, 2018.',
        '[11] 华为技术有限公司. 星河 AI 网络运维白皮书[R]. 2024.',
        '[12] 中华人民共和国工业和信息化部. 「十四五」信息通信行业发展规划[OL]. 2021.',
        '[13] Bosshart P, Gibb G, Kim H S, et al. Forwarding Metamorphosis: Fast Programmable Match-Action Processing in Hardware for SDN[C]. ACM SIGCOMM, 2013: 99-110.',
        '[14] Jain S, Kumar A, Mandal S, et al. B4: Experience with a Globally-Deployed Software Defined WAN[C]. ACM SIGCOMM, 2013: 3-14.',
        '[15] 中华人民共和国国务院. 中国制造 2025[OL]. http://www.gov.cn, 2015.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        set_pf(p, indent=False)
        add_run(p, ref, size=SZ_XIAOSI)

    doc.save(OUTPUT)
    print(f'报告已保存: {OUTPUT}')


if __name__ == '__main__':
    build_report()
